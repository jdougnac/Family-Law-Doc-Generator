from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Length
from docx.text.run import Font
from docx.table import Table
from tkinter import *
from tkinter.filedialog import *
from imr import ingresoMinimo
from numeralizar import *
from decimal import *
import json


def tenerIMR(alimentos, IMR):
    g=Decimal((int(alimentos)/int(IMR))*100)
    if g < 100:
        final = int(g)
        return('l '+numeralizar(str(final))+' por ciento de un Ingreso Mínimo Remuneracional')
    elif g == 100:
        final = 1
        return(' un Ingreso Mínimo Remuneracional')
    elif g > 100:
        final = str(round(g/100,1))
        if final[-2:] == '.0':
            final = int(final)
            return(' ' + tenerNumero(str(final))+ ' Ingresos Mínimos Remuneracionales') 
        elif final[-2:] == '.1':
            final = int(final)
            return(tenerNumero(str(final))+ ' coma un Ingresos Mínimos Remuneracionales')         
        elif final[-2:] == '.2':
            final = int(final)
            return(' ' + tenerNumero(str(final))+ ' coma dos Ingresos Mínimos Remuneracionales')         
        elif final[-2:] == '.3':
            final = int(final)
            return(' ' + tenerNumero(str(final))+ ' coma tres Ingresos Mínimos Remuneracionales')         
        elif final[-2:] == '.4':
            final = int(final)
            return(' ' + tenerNumero(str(final))+ ' coma cuatro Ingresos Mínimos Remuneracionales')         
        elif final[-2:] == '.5':
            final = int(final[:-2])
            return(' ' + tenerNumero(str(final))+ ' coma cinco Ingresos Mínimos Remuneracionales')         
        elif final[-2:] == '.6':
            final = int(final)
            return(' ' + tenerNumero(str(final))+ ' coma seis Ingresos Mínimos Remuneracionales')         
        elif final[-2:] == '.7':
            final = int(final)
            return(' ' + tenerNumero(str(final))+ ' coma siete Ingresos Mínimos Remuneracionales')         
        elif final[-2:] == '.8':
            final = int(final)
            return(' ' + tenerNumero(str(final))+ ' coma ocho Ingresos Mínimos Remuneracionales')         
        elif final[-2:] == '.9':
            final = int(final)
            return(' ' + tenerNumero(str(final))+ ' coma nueve Ingresos Mínimos Remuneracionales')           


def acuerdoCompletoEscPub():
    with open(filedialog.askopenfilename()) as fp:
            data_load=json.load(fp)
           
    
    data_load['hijosMenoresDeEdad']=0
    data_load['listaHijosMenoresDeEdad']=[]
    listaCuidadoPersonal=[]
    listaPatriaPotestad=[]
    
    if data_load['edadDeHijoEntreParte1'] !='' and int(data_load['edadDeHijoEntreParte1'])<18:
        data_load['hijosMenoresDeEdad']+=1
        data_load['listaHijosMenoresDeEdad'].append(data_load['nombreDeHijoEntreParte1'])
        listaCuidadoPersonal.append(data_load['cuidadoPersonalHijo1'])
        listaPatriaPotestad.append(data_load['patriaPotestadHijo1'])
    if data_load['edadDeHijoEntreParte2'] !='' and int(data_load['edadDeHijoEntreParte2'])<18:
        data_load['hijosMenoresDeEdad']+=1
        data_load['listaHijosMenoresDeEdad'].append(data_load['nombreDeHijoEntreParte2'])
        listaCuidadoPersonal.append(data_load['cuidadoPersonalHijo2'])
        listaPatriaPotestad.append(data_load['patriaPotestadHijo2'])
    if data_load['edadDeHijoEntreParte3'] !='' and int(data_load['edadDeHijoEntreParte3'])<18:
        data_load['hijosMenoresDeEdad']+=1
        data_load['listaHijosMenoresDeEdad'].append(data_load['nombreDeHijoEntreParte3'])
        listaCuidadoPersonal.append(data_load['cuidadoPersonalHijo3'])
        listaPatriaPotestad.append(data_load['patriaPotestadHijo3'])
    if data_load['edadDeHijoEntreParte4'] !='' and int(data_load['edadDeHijoEntreParte4'])<18:
        data_load['hijosMenoresDeEdad']+=1
        data_load['listaHijosMenoresDeEdad'].append(data_load['nombreDeHijoEntreParte4'])
        listaCuidadoPersonal.append(data_load['cuidadoPersonalHijo4'])
        listaPatriaPotestad.append(data_load['patriaPotestadHijo4'])
    if data_load['edadDeHijoEntreParte5'] !='' and int(data_load['edadDeHijoEntreParte5'])<18:
        data_load['hijosMenoresDeEdad']+=1        
        data_load['listaHijosMenoresDeEdad'].append(data_load['nombreDeHijoEntreParte5'])
        listaCuidadoPersonal.append(data_load['cuidadoPersonalHijo5'])
        listaPatriaPotestad.append(data_load['patriaPotestadHijo5'])

 



        
    #DOCUMENTO PROPIAMENTE TAL

    document=Document("documentoBase.docx")
    Font.name = 'Century Gothic'
    Font.size = Pt(16)

    parrafo1=document.add_paragraph()
    parrafo1.add_run('                                                                 Rep. Nº ')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_paragraph('')    
    encabezado1=document.add_paragraph()
    encabezado1.add_run('ACUERDO COMPLETO Y SUFICIENTE POR DIVORCIO DE MUTUO ACUERDO').bold=True
    parrafo1=document.add_paragraph()
    parrafo1.add_run('- Y - ').bold=True
    parrafo1=document.add_paragraph()    
    if data_load['modoPagoCompensacion'] == 'Bien Inmueble':
        parrafo1.add_run('CESIÓN DE INMUEBLE POR COMPENSACIÓN ECONÓMICA.')
    elif data_load['modoPagoCompensacion'] == 'Derechos sobre un inmueble' or data_load['modoPagoCompensacion'] == 'Derecho sobre inmueble a liquidar':
        parrafo1.add_run('CESIÓN DE INMUEBLE POR COMPENSACIÓN ECONÓMICA.')
    parrafo1=document.add_paragraph()
    parrafo1.add_run('                     ENTRE')
    parrafo1=document.add_paragraph()
    parrafo1.add_run('                     ')
    parrafo1.add_run(data_load['nombrePatrocinado'].upper())
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoPatrocinado'].upper())
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoPatrocinado'].upper())

    parrafo1=document.add_paragraph()
    parrafo1.add_run('                     -Y-')
    parrafo1=document.add_paragraph()    
    parrafo1.add_run(data_load['nombreContraparte'].upper())
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoContraparte'].upper())
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoContraparte'].upper())
    
    document.add_page_break()    
    
    parrafo1=document.add_paragraph()
    parrafo1.add_run(" Comparecen: ")
    parrafo1.add_run(data_load['honorificoPatrocinado'])    
    parrafo1.add_run(" ")    
    parrafo1.add_run(data_load['nombrePatrocinado'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoPatrocinado'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoPatrocinado'].upper()).bold=True
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['nacionalidadPatrocinado'].lower())    
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['estadoCivilPatrocinado'].lower())
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['profesionPatrocinado'].lower())    
    parrafo1.add_run(", cédula nacional de identidad número ")
    parrafo1.add_run(getRut(data_load['rutPatrocinado']))
    parrafo1.add_run(", con domicilio en calle ")
    parrafo1.add_run(data_load['domicilioPatrocinado'])
    parrafo1.add_run(" número ")
    parrafo1.add_run(tenerNumero(data_load['numeroDomicilioPatrocinado']))
    parrafo1.add_run(", comuna de ")
    parrafo1.add_run(data_load['comunaPatrocinado'])
    parrafo1.add_run(", ciudad de ")
    parrafo1.add_run(data_load['ciudadPatrocinado'])
    parrafo1.add_run(", Región ")    
    parrafo1.add_run(regionPatrocinado)
    parrafo1.add_run(", en adelante como ")
    parrafo1.add_run('"').bold=True
    parrafo1.add_run(data_load['conectorPatrocinado']).bold=True
    parrafo1.add_run(" solicitante o ").bold=True
    parrafo1.add_run(data_load['conectorPatrocinado']).bold=True
    parrafo1.add_run(' cedente"').bold=True
    parrafo1.add_run("; y ")

    parrafo1.add_run(data_load['honorificoContraparte'])
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['nombreContraparte'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoContraparte'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoContraparte'].upper()).bold=True
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['nacionalidadContraparte'].lower())    
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['estadoCivilContraparte'].lower())
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['profesionContraparte'].lower())
    parrafo1.add_run(", cédula nacional de identidad número ")
    parrafo1.add_run(getRut(data_load['rutContraparte']))
    parrafo1.add_run(", con domicilio en calle ")
    parrafo1.add_run(data_load['domicilioContraparte'])
    parrafo1.add_run(" número ")
    parrafo1.add_run(tenerNumero(data_load['numeroDomicilioContraparte']))
    parrafo1.add_run(", comuna de ")
    parrafo1.add_run(data_load['comunaContraparte'])
    parrafo1.add_run(", ciudad de ")
    parrafo1.add_run(data_load['ciudadContraparte'])
    parrafo1.add_run(", Región ")    
    parrafo1.add_run(regionContraparte)
    parrafo1.add_run(", en adelante como ")
    parrafo1.add_run('"').bold=True
    parrafo1.add_run(data_load['conectorContraparte']).bold=True
    parrafo1.add_run(" solicitad").bold=True
    parrafo1.add_run(data_load['generoContraparte']).bold=True
    parrafo1.add_run(" o ").bold=True
    parrafo1.add_run(data_load['conectorContraparte']).bold=True
    parrafo1.add_run(' cesionari').bold=True
    parrafo1.add_run(data_load['generoContraparte']).bold=True
    parrafo1.add_run('"').bold=True    
    parrafo1.add_run("; solicitante y solicitad")
    parrafo1.add_run(data_load['generoContraparte'])    
    parrafo1.add_run(" también como ")
    parrafo1.add_run('"los solicitantes", ').bold=True
    parrafo1.add_run('todos los comparecientes mayores de edad, quienes acreditan su identidad con las cédulas antes citadas, y exponen: ')
    parrafo1.add_run('PRIMERO: ').bold=True
    parrafo1.add_run('Las partes hacen constar lo siguiente: ')
    parrafo1.add_run('A) ').bold=True
    parrafo1.add_run(data_load['honorificoPatrocinado'])    
    parrafo1.add_run(" ")    
    parrafo1.add_run(data_load['nombrePatrocinado'])
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoPatrocinado'])
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoPatrocinado'])
    parrafo1.add_run(' y ')
    parrafo1.add_run(data_load['honorificoContraparte'])
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['nombreContraparte'])
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoContraparte'])
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoContraparte']) 
    parrafo1.add_run(', ambos ya individualizados, son cónyuges, contrajeron matrimonio el día ')
    if data_load['diaDeMatrimonio']=="1":
        parrafo1.add_run("primero")
    else:
        parrafo1.add_run(numeralizar(data_load['diaDeMatrimonio']))
    parrafo1.add_run(" de ")
    parrafo1.add_run(data_load['mesDeMatrimonio'])
    parrafo1.add_run(" de ")
    parrafo1.add_run(numeralizar(data_load['añoDeMatrimonio']))
    parrafo1.add_run(" en la circunscripción de ")
    parrafo1.add_run(data_load['comunaDeMatrimonio'])
    parrafo1.add_run(", con número de inscripción ")
    parrafo1.add_run(numeralizar(data_load['numeroInscripcionMatrimonio']))
    parrafo1.add_run(" del año ")
    parrafo1.add_run(numeralizar(data_load['añoDeMatrimonio']))    
    parrafo1.add_run(", bajo el régimen de ")    
    parrafo1.add_run(data_load['regimenMatrimonial'].lower())
    parrafo1.add_run(".")
    parrafo1.add_run(' B) ').bold=True
    parrafo1.add_run('Que de dicha relación matrimonial ')

    if data_load['hijosEntrePartes']==0:
        parrafo1.add_run(" no nacieron hijos comunes.")
    elif data_load['hijosEntrePartes']==1:
        if 'F' in data_load['sexoDeHijoEntreParte1']:
            parrafo1.add_run(" nació su hija ")
        elif 'M' in data_load['sexoDeHijoEntreParte1']:
            parrafo1.add_run(" nació su hijo ")
        parrafo1.add_run(data_load['nombreDeHijoEntreParte1'])
        parrafo1.add_run(", cédula de identidad número ")
        parrafo1.add_run(getRut(data_load['rutDeHijoEntreParte1']))
        parrafo1.add_run(", de actuales ")
        parrafo1.add_run(numeralizar(data_load['edadDeHijoEntreParte1']))
        parrafo1.add_run(" años de edad. ")
    elif data_load['hijosEntrePartes']>1:                  
        if data_load['hijosEntrePartes'] == 2:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2']:
                parrafo1.add_run("nacieron sus hijas: ")
            else:
                parrafo1.add_run("nacieron sus hijos: ")
        elif data_load['hijosEntrePartes'] == 3:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3']:
                parrafo1.add_run("nacieron sus hijas: ")
            else:
                parrafo1.add_run("nacieron sus hijos: ")
        elif data_load['hijosEntrePartes'] == 4:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3'] and 'F' in data_load['sexoDeHijoEntreParte4']:
                parrafo1.add_run("nacieron sus hijas: ")
            else:
                parrafo1.add_run("nacieron sus hijos: ")            
        elif data_load['hijosEntrePartes'] == 5:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3'] and 'F' in data_load['sexoDeHijoEntreParte4'] and 'F' in data_load['sexoDeHijoEntreParte5']:
                parrafo1.add_run("nacieron sus hijas: ")
            else:
                parrafo1.add_run("nacieron sus hijos: ")            
        parrafo1.add_run(data_load['nombreDeHijoEntreParte1'])
        parrafo1.add_run(", cédula de identidad número ")
        parrafo1.add_run(getRut(data_load['rutDeHijoEntreParte1']))
        parrafo1.add_run(", de actuales ")
        parrafo1.add_run(numeralizar(data_load['edadDeHijoEntreParte1']))
        if data_load['edadDeHijoEntreParte1']== "1":
            parrafo1.add_run(" año de edad; ")
        else:
            parrafo1.add_run(" años de edad; ")
        parrafo1.add_run(data_load['nombreDeHijoEntreParte2'])
        parrafo1.add_run(", cédula de identidad número ")
        parrafo1.add_run(getRut(data_load['rutDeHijoEntreParte2']))
        parrafo1.add_run(", de actuales ")
        parrafo1.add_run(numeralizar(data_load['edadDeHijoEntreParte2']))
        if data_load['edadDeHijoEntreParte2']== "1":
            parrafo1.add_run(" año de edad; ")
        else:
            parrafo1.add_run(" años de edad; ")
        if 'a' in data_load['nombreDeHijoEntreParte3'] or 'e' in data_load['nombreDeHijoEntreParte3'] or 'i' in data_load['nombreDeHijoEntreParte3'] or 'o' in data_load['nombreDeHijoEntreParte3'] or 'u' in data_load['nombreDeHijoEntreParte3']:
            parrafo1.add_run(data_load['nombreDeHijoEntreParte3'])
            parrafo1.add_run(", cédula de identidad número ")
            parrafo1.add_run(getRut(data_load['rutDeHijoEntreParte3']))
            parrafo1.add_run(", de actuales ")
            parrafo1.add_run(numeralizar(data_load['edadDeHijoEntreParte3']))
            if data_load['edadDeHijoEntreParte3']== "1":
                parrafo1.add_run(" año de edad; ")
            else:
                parrafo1.add_run(" años de edad; ")
        if 'a' in data_load['nombreDeHijoEntreParte4'] or 'e' in data_load['nombreDeHijoEntreParte4'] or 'i' in data_load['nombreDeHijoEntreParte4'] or 'o' in data_load['nombreDeHijoEntreParte4'] or 'u' in data_load['nombreDeHijoEntreParte4']:
            parrafo1.add_run(data_load['nombreDeHijoEntreParte4'])
            parrafo1.add_run(", cédula de identidad número ")
            parrafo1.add_run(getRut(data_load['rutDeHijoEntreParte4']))
            parrafo1.add_run(", de actuales ")
            parrafo1.add_run(numeralizar(data_load['edadDeHijoEntreParte4']))
            if data_load['edadDeHijoEntreParte4']== "1":
                parrafo1.add_run(" año de edad; ")
            else:
                parrafo1.add_run(" años de edad; ")

        if 'a' in data_load['nombreDeHijoEntreParte5'] or 'e' in data_load['nombreDeHijoEntreParte5'] or 'i' in data_load['nombreDeHijoEntreParte5'] or 'o' in data_load['nombreDeHijoEntreParte5'] or 'u' in data_load['nombreDeHijoEntreParte5']:
            parrafo1.add_run(data_load['nombreDeHijoEntreParte5'])
            parrafo1.add_run(", cédula de identidad número ")
            parrafo1.add_run(getRut(data_load['rutDeHijoEntreParte5']))
            parrafo1.add_run(", de actuales ")
            parrafo1.add_run(numeralizar(data_load['edadDeHijoEntreParte5']))
            if data_load['edadDeHijoEntreParte5']== "1":
                parrafo1.add_run(" año de edad; ")
            else:
                parrafo1.add_run(" años de edad; ")
                      
    parrafo1.add_run('SEGUNDO: ').bold=True
    parrafo1.add_run('Que los solicitantes acompañan el siguiente acuerdo para que S.S. lo califique de completo y suficiente, autorice y apruebe en la respectiva audiencia, en esta causa de divorcio por cese de la convivencia de común acuerdo. ')    
    parrafo1.add_run('En efecto, ambos solicitantes han acordado regular todas las materias relevantes y que la ley número diecinueve mil novecientos cuarenta y siete, en su artículo cincuenta y cinco, en relación con el artículo veintiuno del mismo cuerpo normativo exige, para que S.S. acoja la presente demanda de divorcio por cese de la convivencia de común acuerdo en la que se solicita que se declare el divorcio respecto del matrimonio celebrado por ambos solicitantes. ')
    parrafo1.add_run('TERCERO: ').bold=True
    parrafo1.add_run('Las materias acordadas son las siguientes: ')
    parrafo1.add_run('A) ').bold=True
    parrafo1.add_run('ALIMENTOS MENORES: ')
    g=ingresoMinimo(data_load['cuantiaAlimentosMenores'],data_load['cuantiaIMR'])
    if data_load['huboMediacion']=='Si':
        parrafo1.add_run(' En proceso de mediación RIT ')
        parrafo1.add_run(tenerRIT(data_load['ritMediacion']))
        parrafo1.add_run(', seguida ante ')
        parrafo1.add_run('VACÍO').bold=True
        parrafo1.add_run(', las partes estipulan que ')
        if data_load['pagaAlimentosMenores']=='Patrocinado':
            parrafo1.add_run(data_load['honorificoPatrocinado'])
            parrafo1.add_run(' ')
            parrafo1.add_run(data_load['nombrePatrocinado'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoPaternoPatrocinado'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoMaternoPatrocinado'])
            parrafo1.add_run(' pagará la suma de ')
            parrafo1.add_run(numeralizar(data_load['cuantiaAlimentosMenores']))
            parrafo1.add_run(' pesos, equivalente a')
            parrafo1.add_run(tenerIMR(data_load['cuantiaAlimentosMenores'],data_load['cuantiaIMR']))
            parrafo1.add_run('.')
        elif data_load['pagaAlimentosMenores']=='Contraparte':
            parrafo1.add_run(data_load['honorificoContraparte'])
            parrafo1.add_run(' ')
            parrafo1.add_run(data_load['nombreContraparte'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoPaternoContraparte'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoMaternoContraparte'])
            parrafo1.add_run(' pagará la suma de ')
            parrafo1.add_run(numeralizar(data_load['cuantiaAlimentosMenores']))
            parrafo1.add_run(' pesos, equivalente a')
            parrafo1.add_run(tenerIMR(data_load['cuantiaAlimentosMenores'],data_load['cuantiaIMR']))
            parrafo1.add_run('.')
        elif data_load['pagaAlimentosMenores']=='No hay':
            parrafo1.add_run('no se pagarán alimentos menores, debido a')
            if data_load['hijosEntrePartes']==0:
                parrafo1.add_run(' que las partes no tienen hijos en común.')
            elif data_load['hijosMenoresDeEdad']==0:
                parrafo1.add_run(' que no hay hijos menores de edad nacidos entre las partes.')
            else:
                parrafo1.add_run('que ')
                parrafo1.add_run('VACÍO').bold=True
    elif data_load['huboMediacion']=='No': 
        parrafo1.add_run(' Las partes estipulan que ')
        if data_load['pagaAlimentosMenores']=='No hay':
            parrafo1.add_run('no se pagarán alimentos menores, debido a')
            if data_load['hijosEntrePartes']==0:
                parrafo1.add_run(' que las partes no tienen hijos en común.')
            elif data_load['hijosMenoresDeEdad']==0:
                parrafo1.add_run(' que no hay hijos menores de edad nacidos de las partes.')
            else:
                parrafo1.add_run('que ')
                parrafo1.add_run('VACÍO').bold=True          
        else:
            if data_load['pagaAlimentosMenores']=='Patrocinado':
                parrafo1.add_run(data_load['honorificoPatrocinado'])
                parrafo1.add_run(' ')
                parrafo1.add_run(data_load['nombrePatrocinado'])
                parrafo1.add_run(" ")
                parrafo1.add_run(data_load['apellidoPaternoPatrocinado'])
                parrafo1.add_run(" ")
                parrafo1.add_run(data_load['apellidoMaternoPatrocinado'])
                parrafo1.add_run(' pagará la suma de $')
                parrafo1.add_run(numeralizar(data_load['cuantiaAlimentosMenores']))
                parrafo1.add_run(' pesos, equivalente a')
                parrafo1.add_run(tenerIMR(data_load['cuantiaAlimentosMenores'],data_load['cuantiaIMR']))
                parrafo1.add_run('.')
            elif data_load['pagaAlimentosMenores']=='Contraparte':
                parrafo1.add_run(data_load['honorificoContraparte'])
                parrafo1.add_run(' ')
                parrafo1.add_run(data_load['nombreContraparte'])
                parrafo1.add_run(" ")
                parrafo1.add_run(data_load['apellidoPaternoContraparte'])
                parrafo1.add_run(" ")
                parrafo1.add_run(data_load['apellidoMaternoContraparte'])
                parrafo1.add_run(' pagará la suma de $')
                parrafo1.add_run(numeralizar(data_load['cuantiaAlimentosMenores']))
                parrafo1.add_run(' pesos, equivalente a')
                parrafo1.add_run(tenerIMR(data_load['cuantiaAlimentosMenores'],data_load['cuantiaIMR']))
                parrafo1.add_run('.')    


    parrafo1.add_run(' B) ').bold=True
    parrafo1.add_run('ALIMENTOS MAYORES: ')               
    parrafo1.add_run('No se han demandado alimentos entre los cónyuges, por lo que no existen deudas por este concepto.')
    parrafo1.add_run(' C) ').bold=True                
    parrafo1.add_run('PATRIA POTESTAD: ') 
    if int(data_load['hijosEntrePartes'] )==0:
        parrafo1.add_run('No es necesario regular por no haber hijos entre las partes.')
    elif int(data_load['hijosEntrePartes'] )>0:
        if data_load['hijosMenoresDeEdad']==0:
            parrafo1.add_run('No es necesario regular por no haber hijos menores de edad.')
        elif data_load['hijosMenoresDeEdad']>0:
            for x in range(0,len(data_load['listaHijosMenoresDeEdad'])):
                parrafo1.add_run('En el caso de ')
                parrafo1.add_run(data_load['listaHijosMenoresDeEdad'][x])
                parrafo1.add_run(', su cuidado personal corresponderá a ')
                if listaPatriaPotestad[x]=='Patrocinado':
                    parrafo1.add_run(data_load['honorificoPatrocinado'])
                    parrafo1.add_run(' ')
                    parrafo1.add_run(data_load['nombrePatrocinado'])
                    parrafo1.add_run(" ")
                    parrafo1.add_run(data_load['apellidoPaternoPatrocinado'])
                    parrafo1.add_run(" ")
                    parrafo1.add_run(data_load['apellidoMaternoPatrocinado'])
                    parrafo1.add_run('. ')
                elif listaPatriaPotestad[x]=='Contraparte':
                    parrafo1.add_run(data_load['honorificoContraparte'])
                    parrafo1.add_run(' ')
                    parrafo1.add_run(data_load['nombreContraparte'])
                    parrafo1.add_run(" ")
                    parrafo1.add_run(data_load['apellidoPaternoContraparte'])
                    parrafo1.add_run(" ")
                    parrafo1.add_run(data_load['apellidoMaternoContraparte'])
                    parrafo1.add_run('. ') 
    parrafo1.add_run('D) ').bold=True                
    parrafo1.add_run('CUIDADO PERSONAL: ')
    if int(data_load['hijosEntrePartes'] )==0:
        parrafo1.add_run(' No corresponde por no haber hijos entre las partes. ')
    elif int(data_load['hijosEntrePartes'] )>0:
        if data_load['hijosMenoresDeEdad']==0:
            parrafo1.add_run(' No corresponde por no haber hijos menores de edad. ')
        elif data_load['hijosMenoresDeEdad']>0:
            for x in range(0,len(data_load['listaHijosMenoresDeEdad'])):
                parrafo1.add_run('En el caso de ')
                parrafo1.add_run(data_load['listaHijosMenoresDeEdad'][x])
                parrafo1.add_run(', su cuidado personal corresponderá a ')
                if listaCuidadoPersonal[x]=='Patrocinado':
                    parrafo1.add_run(data_load['honorificoPatrocinado'])
                    parrafo1.add_run(' ')
                    parrafo1.add_run(data_load['nombrePatrocinado'])
                    parrafo1.add_run(" ")
                    parrafo1.add_run(data_load['apellidoPaternoPatrocinado'])
                    parrafo1.add_run(" ")
                    parrafo1.add_run(data_load['apellidoMaternoPatrocinado'])
                    parrafo1.add_run('. ')
                elif listaCuidadoPersonal[x]=='Contraparte':
                    parrafo1.add_run(data_load['honorificoContraparte'])
                    parrafo1.add_run(' ')
                    parrafo1.add_run(data_load['nombreContraparte'])
                    parrafo1.add_run(" ")
                    parrafo1.add_run(data_load['apellidoPaternoContraparte'])
                    parrafo1.add_run(" ")
                    parrafo1.add_run(data_load['apellidoMaternoContraparte'])
                    parrafo1.add_run('. ')    

    parrafo1.add_run('E) ').bold=True  
    parrafo1.add_run('RELACIÓN DIRECTA Y REGULAR: ')
    if int(data_load['hijosEntrePartes'] )==0:
        parrafo1.add_run(' No corresponde por no haber hijos entre las partes.')
    elif int(data_load['hijosEntrePartes'] )>0:
        if data_load['hijosMenoresDeEdad']==0:
            parrafo1.add_run(' No corresponde por no haber hijos menores de edad. ')
        elif data_load['hijosMenoresDeEdad']>0:
            for x in range(0,len(data_load['listaHijosMenoresDeEdad'])):
                parrafo1.add_run('En el caso de ')
                parrafo1.add_run(data_load['listaHijosMenoresDeEdad'][x])
                parrafo1.add_run(', los solicitantes han acordado que ')  
                parrafo1.add_run('VACÍO. ').bold=True
    else:
        parrafo1.add_run('A este respecto se establece lo siguiente:')
        parrafo1.add_run(' VACÍO. ').bold=True
    parrafo1.add_run('F) ').bold=True
    parrafo1.add_run(' RÉGIMEN PATRIMONIAL:')
    parrafo1.add_run(' El régimen patrimonial pactado al momento de celebración del matrimonio fue el de ')
    parrafo1.add_run(data_load['regimenMatrimonial'])
    parrafo1.add_run('.')
    
    parrafo1.add_run('F) ').bold=True          
    parrafo1.add_run('COMPENSACIÓN ECONÓMICA:')    
    if 'no' in data_load['pideCompensacionEconomica']:
        parrafo1.add_run('Habiendo sido informados sobre la institución de la compensación económica, ambas partes renuncian expresamente a ella')
    elif data_load['pideCompensacionEconomica']=='Patrocinado':
        parrafo1.add_run(' Las partes acuerdan que ')
        parrafo1.add_run(data_load['honorificoContraparte'])
        parrafo1.add_run(' ')
        parrafo1.add_run(data_load['nombreContraparte'])
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoPaternoContraparte'])
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoMaternoContraparte'])  
        parrafo1.add_run(' se obliga a pagar a ')  
        parrafo1.add_run(data_load['honorificoPatrocinado'])
        parrafo1.add_run(' ')
        parrafo1.add_run(data_load['nombrePatrocinado'])
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoPaternoPatrocinado'])
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoMaternoPatrocinado'])         
    elif data_load['pideCompensacionEconomica']=='Contraparte':
        parrafo1.add_run(' Las partes acuerdan que ')
        parrafo1.add_run(data_load['honorificoPatrocinado'])
        parrafo1.add_run(' ')
        parrafo1.add_run(data_load['nombrePatrocinado'])
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoPaternoPatrocinado'])
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoMaternoPatrocinado'])  
        parrafo1.add_run(' se obliga a pagar a ')  
        parrafo1.add_run(data_load['honorificoContraparte'])
        parrafo1.add_run(' ')
        parrafo1.add_run(data_load['nombreContraparte'])
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoPaternoContraparte'])
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoMaternoContraparte']) 
    parrafo1.add_run(', a título de compensación económica, el valor de ')
    parrafo1.add_run(numeralizar(data_load['valorCompensacion']))
    zz=numeralizar(data_load['valorCompensacion'])
    if zz[-8:] == "millones":
        parrafo1.add_run(" de")
    parrafo1.add_run(' pesos. Éste se hará efectivo')
    if data_load['modoPagoCompensacion'] == 'Bien Inmueble':
        parrafo1.add_run(', como pago equivalente a dicho monto, mediante la cesión del inmueble ')
        parrafo1.add_run(' ubicado en ')
        parrafo1.add_run(data_load['direccionInmueblePagoCompensacion'])
        parrafo1.add_run(' número ')
        parrafo1.add_run(tenerNumero(data_load['numeroInmueblePagoCompensacion']))
        parrafo1.add_run(' de propiedad de ')
        if data_load['pideCompensacionEconomica']=='Contraparte':
            parrafo1.add_run(data_load['honorificoPatrocinado'])
            parrafo1.add_run(' ')
            parrafo1.add_run(data_load['nombrePatrocinado'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoPaternoPatrocinado'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoMaternoPatrocinado']) 
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo1.add_run(data_load['honorificoContraparte'])
            parrafo1.add_run(' ')
            parrafo1.add_run(data_load['nombreContraparte'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoPaternoContraparte'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoMaternoContraparte']) 
        parrafo1.add_run('.')

    elif data_load['modoPagoCompensacion'] == 'Derechos sobre un inmueble':
        parrafo1.add_run(', como pago equivalente a dicho monto, mediante la cesión de los derechos que ')
        if data_load['pideCompensacionEconomica']=='Contraparte':
            parrafo1.add_run(data_load['honorificoPatrocinado'])
            parrafo1.add_run(' ')
            parrafo1.add_run(data_load['nombrePatrocinado'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoPaternoPatrocinado'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoMaternoPatrocinado']) 
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo1.add_run(data_load['honorificoContraparte'])
            parrafo1.add_run(' ')
            parrafo1.add_run(data_load['nombreContraparte'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoPaternoContraparte'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoMaternoContraparte']) 
        parrafo1.add_run(' tiene sobre el inmueble ubicado en ')
        parrafo1.add_run(data_load['direccionInmueblePagoCompensacion'])
        parrafo1.add_run(' número ')
        parrafo1.add_run(tenerNumero(data_load['numeroInmueblePagoCompensacion']))
        parrafo1.add_run('.') 
    elif data_load['modoPagoCompensacion'] == 'Derechos sobre inmueble sociedad conyugal': 
        parrafo1.add_run(', como pago equivalente a dicho monto, mediante la cesión de los derechos que a ')
        if data_load['pideCompensacionEconomica']=='Contraparte':
            parrafo1.add_run(data_load['honorificoPatrocinado'])
            parrafo1.add_run(' ')
            parrafo1.add_run(data_load['nombrePatrocinado'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoPaternoPatrocinado'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoMaternoPatrocinado']) 
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo1.add_run(data_load['honorificoContraparte'])
            parrafo1.add_run(' ')
            parrafo1.add_run(data_load['nombreContraparte'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoPaternoContraparte'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoMaternoContraparte']) 
        parrafo1.add_run(' corresponden sobre el inmueble propiedad de la sociedad conyugal ubicado en ')
        parrafo1.add_run(data_load['direccionInmueblePagoCompensacion'])
        parrafo1.add_run(' número ')
        parrafo1.add_run(tenerNumero(data_load['numeroInmueblePagoCompensacion']))
        parrafo1.add_run('.')     
    elif data_load['modoPagoCompensacion'] == 'Dinero':
        parrafo1.add_run(' mediante el pago de la suma de ')
        parrafo1.add_run(numeralizar(data_load['valorCompensacion']))
        if zz[-8:] == "millones":
            parrafo1.add_run(" de")        
        parrafo1.add_run(' pesos.')
    elif data_load['modoPagoCompensacion'] == 'Otro': 
        parrafo1.add_run(' mediante')    
        parrafo1.add_run(' VACÍO.').bold=True
        
    parrafo1.add_run(' CUARTO:').bold=True
    parrafo1.add_run(' que en definitiva ')
    if data_load['pideCompensacionEconomica']=='Contraparte':
        parrafo1.add_run(data_load['honorificoPatrocinado'].upper()).bold=True
        parrafo1.add_run(' ')
        parrafo1.add_run(data_load['nombrePatrocinado'].upper()).bold=True
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoPaternoPatrocinado'].upper()).bold=True
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoMaternoPatrocinado'].upper()).bold=True
        parrafo1.add_run(' cede y transfiere a ')
        parrafo1.add_run(data_load['honorificoContraparte'].upper()).bold=True
        parrafo1.add_run(' ')
        parrafo1.add_run(data_load['nombreContraparte'].upper()).bold=True
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoPaternoContraparte'].upper()).bold=True
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoMaternoContraparte'].upper()).bold=True        
    elif data_load['pideCompensacionEconomica']=='Patrocinado':
        parrafo1.add_run(data_load['honorificoContraparte'].upper()).bold=True
        parrafo1.add_run(' ')
        parrafo1.add_run(data_load['nombreContraparte'].upper()).bold=True
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoPaternoContraparte'].upper()).bold=True
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoMaternoContraparte'].upper()).bold=True
        parrafo1.add_run(' cede y transfiere a ')        
        parrafo1.add_run(data_load['honorificoPatrocinado'].upper()).bold=True
        parrafo1.add_run(' ')
        parrafo1.add_run(data_load['nombrePatrocinado'].upper()).bold=True
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoPaternoPatrocinado'].upper()).bold=True
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoMaternoPatrocinado'].upper()).bold=True
    if data_load['modoPagoCompensacion']=='Bien Inmueble':
        parrafo1.add_run(', quien acepta y adquiere para sí, como únic')
        if data_load['pideCompensacionEconomica']=='Contraparte': 
            parrafo1.add_run(data_load['generoContraparte'])
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo1.add_run(data_load['generoPatrocinado'])
        parrafo1.add_run(' y exclusiv')
        if data_load['pideCompensacionEconomica']=='Contraparte': 
            parrafo1.add_run(data_load['generoContraparte'])
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo1.add_run(data_load['generoPatrocinado'])
        parrafo1.add_run(' dueñ')
        if data_load['pideCompensacionEconomica']=='Contraparte': 
            parrafo1.add_run(data_load['generoContraparte'])
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo1.add_run(data_load['generoPatrocinado'])
        parrafo1.add_run(' de la propiedad ubicada en ')    
        parrafo1.add_run(data_load['direccionInmueblePagoCompensacion'])
        parrafo1.add_run(' número ')
        parrafo1.add_run(tenerNumero(data_load['numeroInmueblePagoCompensacion']))
        parrafo1.add_run('.')         
            
    elif data_load['modoPagoCompensacion']== 'Derechos sobre un inmueble' or data_load['modoPagoCompensacion']== 'Derechos sobre inmueble sociedad conyugal':
        parrafo1.add_run(', quien los acepta y adquiere para sí, quedando así como únic')
        if data_load['pideCompensacionEconomica']=='Contraparte': 
            parrafo1.add_run(data_load['generoContraparte'])
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo1.add_run(data_load['generoPatrocinado'])
        parrafo1.add_run(' y exclusiv')
        if data_load['pideCompensacionEconomica']=='Contraparte': 
            parrafo1.add_run(data_load['generoContraparte'])
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo1.add_run(data_load['generoPatrocinado'])
        parrafo1.add_run(' dueñ')
        if data_load['pideCompensacionEconomica']=='Contraparte': 
            parrafo1.add_run(data_load['generoContraparte'])
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo1.add_run(data_load['generoPatrocinado'])
        parrafo1.add_run(' de la propiedad ubicada en ')    
        parrafo1.add_run(data_load['direccionInmueblePagoCompensacion'])
        parrafo1.add_run(' número ')
        parrafo1.add_run(tenerNumero(data_load['numeroInmueblePagoCompensacion']))
        parrafo1.add_run('.')       
    elif data_load['modoPagoCompensacion']=='Dinero':
        parrafo1.add_run(' la suma de ')
        parrafo1.add_run(numeralizar(data_load['valorCompensacion']))
        if zz[-8:] == "millones":
            parrafo1.add_run(" de")        
        parrafo1.add_run(' pesos.')
    elif data_load['modoPagoCompensacion']=='Otro':
        parrafo1.add_run('VACÍO').bold=True
    formatoParrafo1=parrafo1.paragraph_format
    formatoParrafo1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    formatoParrafo1.line_spacing=1.5        
        
        



    
    document.save('acuerdoEscPub'+data_load['apellidoPaternoPatrocinado']+data_load['apellidoPaternoContraparte']+'.docx')

    print("I am a zombie, ZOMBIE!!")
