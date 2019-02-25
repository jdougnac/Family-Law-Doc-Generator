from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Length
from docx.text.run import Font
from docx.table import Table
from tkinter import *
from tkinter.filedialog import *
import json

def patrocinioYPoderFin():
    with open(filedialog.askopenfilename()) as fp:
            data_load=json.load(fp)

#DOCUMENTO PROPIAMENTE TAL

    document=Document("documentoBase.docx")
    Font.name = 'Century Gothic'
    Font.size = Pt(16)
    
    suma="patrocinio y poder"
    encabezado1=document.add_paragraph()
    encabezado1.add_run(suma.upper()).bold=True


    encabezado2=document.add_paragraph()
    encabezado2.add_run("S.J. DE FAMILIA DE SANTIAGO (").bold=True
    encabezado2.add_run(data_load['juzgado'].upper()).bold=True
    encabezado2.add_run(")")
    formatoEncabezado2=encabezado2.paragraph_format
    formatoEncabezado2.alignment=WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    parrafo1=document.add_paragraph()
    parrafo1.add_run("     ")
    parrafo1.add_run(data_load['nombreAbogadoPatrocinado'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoAbogadoPatrocinado'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoAbogadoPatrocinado'].upper()).bold=True
    parrafo1.add_run(", cedula nacional de identidad numero ")
    parrafo1.add_run(data_load['rutAbogadoPatrocinado'])    
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['profesionAbogadoPatrocinado'])    
    parrafo1.add_run(" por la parte ")
    parrafo1.add_run(data_load['situacionProcesalPatrocinado'])
    parrafo1.add_run(" en autos sobre ")
    parrafo1.add_run(data_load['materiaCausa'].upper()).bold=True
    parrafo1.add_run(", caratulados ")
    parrafo1.add_run(data_load['caratulaCausa'].upper()).bold=True
    parrafo1.add_run(", en causa RIT ")
    parrafo1.add_run(data_load['ritDeCausa']).bold=True
    parrafo1.add_run(", a U.S., respetuosamente digo:")

    formatoParrafo1=parrafo1.paragraph_format
    formatoParrafo1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    formatoParrafo1.line_spacing=1.5


    document.add_paragraph()

    parrafo2=document.add_paragraph()
    parrafo2.add_run("     Que por este acto, pongo en conocimiento de S.S. que renuncio al patrocinio y poder conferido por la solicitante en autos, atendido a que finalizó la gestión judicial para el cual fue conferido. De esta forma, se ha dado cumplimiento a lo dispuesto por el artículo 10 inciso segundo del Código de Procedimiento Civil, que dispone: ")
    parrafo2.add_run('“si la causa de la expiración del mandato es la renuncia del procurador, estará éste obligado a ponerla en conocimiento de su mandante, junto con el estado del juicio, y se entenderá vigente el poder hasta que haya transcurrido el término de emplazamiento desde la notificación de la renuncia al mandante”. ').italic=True

    formatoParrafo2=parrafo2.paragraph_format
    formatoParrafo2.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    formatoParrafo2.line_spacing=1.5

    parrafo3=document.add_paragraph()
    parrafo3.add_run("POR TANTO,").bold=True
    parrafo4=document.add_paragraph()
    parrafo4.add_run("PIDO A US.").bold=True
    parrafo5=document.add_paragraph(" Tener por renunciado patrocinio y poder conferido por la solicitante y ordenar la notificacion correspondiente del mismo.")

    document.save('renuncia'+data_load['apellidoPaternoPatrocinado']+'.docx')
    print("I am a zombie, ZOMBIE!!")   
