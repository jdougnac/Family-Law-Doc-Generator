from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Length
from docx.text.run import Font
from docx.table import Table
from tkinter import *
from tkinter.filedialog import *
from numeralizar import *
import json


def alimentosMenores():
    with open(filedialog.askopenfilename()) as fp:
            data_load=json.load(fp)
            

    document=Document("documentoBase.docx")
    Font.name = 'Century Gothic'
    Font.size = Pt(16)


    tablaDatos=document.add_table(rows=3, cols=2)

    proc1=tablaDatos.cell(0,0)
    proc1.paragraphs[0].add_run('PROCEDIMIENTO').bold = True

    proc11=tablaDatos.cell(0,1)
    proc11.text=': ORDINARIO'


    proc2=tablaDatos.cell(1,0)
    proc2.paragraphs[0].add_run('MATERIA').bold = True
    proc22=tablaDatos.cell(1,1)
    proc22.text=': ALIMENTOS MENORES'
        
    document.add_paragraph('') 

    tablaPatrocinado=document.add_table(rows=3, cols=2)

    proc1=tablaPatrocinado.cell(0,0)
    proc1.paragraphs[0].add_run('DEMANDANTE').bold = True
    proc11=tablaPatrocinado.cell(0,1)
    proc11.paragraphs[0].add_run(': ')
    proc11.paragraphs[0].add_run(data_load['nombrePatrocinado'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoPaternoPatrocinado'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoMaternoPatrocinado'])      

    proc2=tablaPatrocinado.cell(1,0)
    proc2.paragraphs[0].add_run('RUT').bold = True
    proc22=tablaPatrocinado.cell(1,1)
    proc22.paragraphs[0].add_run(': ')
    proc22.paragraphs[0].add_run(data_load['rutPatrocinado'])

    proc3=tablaPatrocinado.cell(2,0)
    proc3.paragraphs[0].add_run('DOMICILIO').bold = True

    proc33=tablaPatrocinado.cell(2,1)
    proc33.paragraphs[0].add_run(': ')    
    proc33.paragraphs[0].add_run(data_load['domicilioPatrocinado'])
    proc33.paragraphs[0].add_run(" ")
    proc33.paragraphs[0].add_run(data_load['numeroDomicilioPatrocinado']) 
    proc33.paragraphs[0].add_run(", comuna de ")
    proc33.paragraphs[0].add_run(data_load['comunaPatrocinado'])       

    document.add_paragraph('')   

    tablaAbogadoPatrocinado=document.add_table(rows=4, cols=2)

    proc1=tablaAbogadoPatrocinado.cell(0,0)
    proc1.paragraphs[0].add_run(data_load['profesionAbogadoPatrocinado'].upper()).bold = True
    proc1.paragraphs[0].add_run(' PATROCINANTE').bold = True
    proc11=tablaAbogadoPatrocinado.cell(0,1)
    proc11.paragraphs[0].add_run(': ')
    proc11.paragraphs[0].add_run(data_load['nombreAbogadoPatrocinado'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoPaternoAbogadoPatrocinado'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoMaternoAbogadoPatrocinado'])      

    proc2=tablaAbogadoPatrocinado.cell(1,0)
    proc2.paragraphs[0].add_run('RUT').bold = True
    proc22=tablaAbogadoPatrocinado.cell(1,1)
    proc22.paragraphs[0].add_run(': ')
    proc22.paragraphs[0].add_run(data_load['rutAbogadoPatrocinado'])

    proc3=tablaAbogadoPatrocinado.cell(2,0)
    proc3.paragraphs[0].add_run('DOMICILIO').bold = True

    proc33=tablaAbogadoPatrocinado.cell(2,1)
    proc33.paragraphs[0].add_run(': ')    
    proc33.paragraphs[0].add_run(data_load['domicilioAbogadoPatrocinado'])
    proc33.paragraphs[0].add_run(" ")    
    proc33.paragraphs[0].add_run(data_load['numeroDomicilioAbogadoPatrocinado']) 
    proc33.paragraphs[0].add_run(", comuna de ")
    proc33.paragraphs[0].add_run(data_load['comunaAbogadoPatrocinado'])      

    proc4=tablaAbogadoPatrocinado.cell(3,0)
    proc4.paragraphs[0].add_run('CORREO ELECTRÓNICO').bold = True

    proc44=tablaAbogadoPatrocinado.cell(3,1)
    proc44.paragraphs[0].add_run(': ')    
    proc44.paragraphs[0].add_run(data_load['emailAbogadoPatrocinado'])

    document.add_paragraph('')    


    if 'a' in data_load['nombreApoderadoPatrocinado'] or 'e' in data_load['nombreApoderadoPatrocinado'] or 'i' in data_load['nombreApoderadoPatrocinado'] or 'o' in data_load['nombreApoderadoPatrocinado'] or 'u' in data_load['nombreApoderadoPatrocinado']:
        tablaApoderadoPatrocinado=document.add_table(rows=4, cols=2)

        proc1=tablaApoderadoPatrocinado.cell(0,0)
        proc1.paragraphs[0].add_run(profesionApoderadoPatrocinado.upper()).bold = True
        proc1.paragraphs[0].add_run(' APODERAD').bold = True
        proc1.paragraphs[0].add_run(data_load['generoApoderadoPatrocinado'].upper()).bold = True        
        proc11=tablaApoderadoPatrocinado.cell(0,1)
        proc11.paragraphs[0].add_run(': ')
        proc11.paragraphs[0].add_run(data_load['nombreApoderadoPatrocinado'])
        proc11.paragraphs[0].add_run(' ')
        proc11.paragraphs[0].add_run(data_load['apellidoPaternoApoderadoPatrocinado'])
        proc11.paragraphs[0].add_run(' ')
        proc11.paragraphs[0].add_run(data_load['apellidoMaternoApoderadoPatrocinado'])      

        proc2=tablaApoderadoPatrocinado.cell(1,0)
        proc2.paragraphs[0].add_run('RUT').bold = True
        proc22=tablaApoderadoPatrocinado.cell(1,1)
        proc22.paragraphs[0].add_run(': ')
        proc22.paragraphs[0].add_run(data_load['rutApoderadoPatrocinado'])

        proc3=tablaApoderadoPatrocinado.cell(2,0)
        proc3.paragraphs[0].add_run('DOMICILIO').bold = True

        proc33=tablaApoderadoPatrocinado.cell(2,1)
        proc33.paragraphs[0].add_run(': ')    
        proc33.paragraphs[0].add_run(data_load['domicilioApoderadoPatrocinado'])
        proc33.paragraphs[0].add_run(" ")    
        proc33.paragraphs[0].add_run(data_load['numeroDomicilioApoderadoPatrocinado']) 
        proc33.paragraphs[0].add_run(", comuna de ")
        proc33.paragraphs[0].add_run(data_load['comunaApoderadoPatrocinado'])           

        proc4=tablaApoderadoPatrocinado.cell(3,0)
        proc4.paragraphs[0].add_run('CORREO ELECTRÓNICO').bold = True

        proc44=tablaApoderadoPatrocinado.cell(3,1)
        proc44.paragraphs[0].add_run(': ')    
        proc44.paragraphs[0].add_run(data_load['emailApoderadoPatrocinado'])      


    document.add_paragraph('')   

    tablaContraparte=document.add_table(rows=3, cols=2)

    proc1=tablaContraparte.cell(0,0)
    proc1.paragraphs[0].add_run('DEMANDAD').bold = True
    proc1.paragraphs[0].add_run(data_load['generoContraparte'].upper()).bold = True
    proc11=tablaContraparte.cell(0,1)
    proc11.paragraphs[0].add_run(': ')
    proc11.paragraphs[0].add_run(data_load['nombreContraparte'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoPaternoContraparte'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoMaternoContraparte'])      

    proc2=tablaContraparte.cell(1,0)
    proc2.paragraphs[0].add_run('RUT').bold = True
    proc22=tablaContraparte.cell(1,1)
    proc22.paragraphs[0].add_run(': ')
    proc22.paragraphs[0].add_run(data_load['rutContraparte'])

    proc3=tablaContraparte.cell(2,0)
    proc3.paragraphs[0].add_run('DOMICILIO').bold = True

    proc33=tablaContraparte.cell(2,1)
    proc33.paragraphs[0].add_run(': ')    
    proc33.paragraphs[0].add_run(data_load['domicilioContraparte'])
    proc33.paragraphs[0].add_run(" ")    
    proc33.paragraphs[0].add_run(data_load['numeroDomicilioContraparte'])
    proc33.paragraphs[0].add_run(", comuna de ")
    proc33.paragraphs[0].add_run(data_load['comunaContraparte'])     

    document.add_paragraph('')   


    resumen1=document.add_paragraph()
    resumen1.add_run("EN LO PRINCIPAL: ").bold=True
    resumen1.add_run("Demanda de alimentos menores; ")
    resumen1.add_run("PRIMER OTROSÍ: ").bold=True
    resumen1.add_run("Solicita alimentos provisorios; ")
    resumen1.add_run("SEGUNDO OTROSÍ: ").bold=True
    resumen1.add_run("Acompaña documentos; ")
    resumen1.add_run("TERCER OTROSÍ: ").bold=True
    resumen1.add_run("Medio de notificación que indica; ")
    resumen1.add_run("CUARTO OTROSÍ: ").bold=True
    resumen1.add_run("Patrocinio y poder; ")

    document.add_paragraph("")

    encabezado2=document.add_paragraph()
    encabezado2.add_run("S.J. DE FAMILIA DE SANTIAGO").bold=True
    formatoEncabezado2=encabezado2.paragraph_format
    formatoEncabezado2.alignment=WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")


    parrafo1=document.add_paragraph()

    parrafo1.add_run("    ")
    parrafo1.add_run(data_load['nombrePatrocinado'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoPatrocinado'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoPatrocinado'].upper()).bold=True
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['estadoCivilPatrocinado'].lower())
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['nacionalidadPatrocinado'].lower())
    parrafo1.add_run(", cédula de identidad Nº ")
    parrafo1.add_run(data_load['rutPatrocinado'])
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['profesionPatrocinado'].lower())
    parrafo1.add_run(", domiciliad")
    parrafo1.add_run(data_load['generoPatrocinado'])
    parrafo1.add_run(" en ")
    parrafo1.add_run(data_load['domicilioPatrocinado'])
    parrafo1.add_run(" número ")
    parrafo1.add_run(data_load['numeroDomicilioPatrocinado'])     
    parrafo1.add_run(", comuna de ")
    parrafo1.add_run(data_load['comunaPatrocinado'])
    parrafo1.add_run(", a S.S. respetuosamente digo:")

    parrafo1=document.add_paragraph()    
    
    parrafo1.add_run(" Que vengo en interponer demanda de alimentos menores en contra de ")
    parrafo1.add_run(data_load['honorificoContraparte'])
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['nombreContraparte'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoContraparte'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoContraparte'].upper()).bold=True
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['estadoCivilContraparte'].lower())
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['nacionalidadContraparte'].lower())
    parrafo1.add_run(", cédula de identidad Nº ")
    parrafo1.add_run(data_load['rutContraparte'])
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['profesionContraparte'].lower())
    parrafo1.add_run(", domiciliad")
    parrafo1.add_run(data_load['generoContraparte'])
    parrafo1.add_run(" en ")
    parrafo1.add_run(data_load['domicilioContraparte'])
    parrafo1.add_run(" número ")
    parrafo1.add_run(data_load['numeroDomicilioContraparte'])     
    parrafo1.add_run(", comuna de ")
    parrafo1.add_run(data_load['comunaContraparte'])
    parrafo1.add_run(", en atención a los fundamentos de hecho y de Derecho que a continuación expongo.")

    formatoParrafo1=parrafo1.paragraph_format
    formatoParrafo1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    formatoParrafo1.line_spacing=1.5

    parrafo1=document.add_paragraph()
    parrafo1.add_run("LOS HECHOS:").bold=True
    document.add_paragraph()    
    
    
    if data_load['mesCeseRelacion']!='' and data_load['mesDeMatrimonio']=='':
        parrafo1=document.add_paragraph()
        parrafo1=document.add_run(" Con ")
        parrafo1=document.add_run(data_load['conectorContraparte'])
        parrafo1=document.add_run(" demandad")
        parrafo1=document.add_run(data_load['generoContraparte'])
        parrafo1=document.add_run(" mantuvimos una relación sentimental que duró ")
        if data_load['añosDuracionRelacion']=="0" or data_load['añosDuracionRelacion']=='':
            if data_load['mesesDuracionRelacion']=="1":
                parrafo1.add_run("un mes.")
            else:
                parrafo1.add_run(data_load['mesesDuracionRelacion'])
                parrafo1.add_run(" meses, ")
        elif data_load['añosDuracionRelacion']=="1":
            parrafo1.add_run(" un año")
            if data_load['mesesDuracionRelacion']=="0" or data_load['mesesDuracionRelacion']=='':
                parrafo1.add_run(", ")
            elif data_load['mesesDuracionRelacion']=="1":
                parrafo1.add_run(" y un mes, ")
            else:
                parrafo1.add_run(" y ")
                parrafo1.add_run(data_load['mesesDuracionRelacion'])
                parrafo1.add_run(" meses, ")            
        parrafo1=document.add_run("la cual terminó de forma definitiva en ")
        parrafo1=document.add_run(data_load['mesCeseRelacion'])
        parrafo1=document.add_run(" del año ")
        parrafo1=document.add_run(data_load['añoCeseRelacion'])
        parrafo1=document.add_run(", producto de la cual ")
        


    if data_load['mesCeseRelacion']=='' and data_load['mesDeMatrimonio']!='':
        parrafo1=document.add_paragraph()
        parrafo1.add_run(" Con fecha ")
        parrafo1.add_run(data_load['diaDeMatrimonio'])
        parrafo1.add_run(' de ')
        parrafo1.add_run(data_load['mesDeMatrimonio'])
        parrafo1.add_run(' de ')
        parrafo1.add_run(data_load['añoDeMatrimonio'])        
        parrafo1.add_run(" contrajimos matrimonio ante Oficial de Registro Civil de la circunscripción de ")
        parrafo1.add_run(data_load['comunaDeMatrimonio'])
        parrafo1.add_run(", el que fue anotado bajo la inscripción Nº ")
        parrafo1.add_run(data_load['numeroInscripcionMatrimonio'])
        parrafo1.add_run(" de dicho año pactando el régimen matrimonial de ")
        parrafo1.add_run(data_load['regimenMatrimonial'].lower())        
        parrafo1.add_run(", el cual no ha sido modificado hasta el día de hoy.")
        parrafo1.add_run(" Nuestra convivencia fue normal aproximadamente por ")
        parrafo1.add_run(str(data_load['duracionMatrimonio']))
        parrafo1.add_run(" años, pero al transcurrir el tiempo empezamos a tener desavenencias e incompatibilidades irreconciliables, lo que hizo que nuestra convivencia fuera insoportable para seguir manteniéndola,  razón por la cual nos separamos de hecho el año ")
        parrafo1.add_run(data_load['añoDeCeseConvivencia'])
        parrafo1.add_run(", fecha en que ")
        if 'atrocinado' in data_load['abandonaHogarComun']:              
            parrafo1.add_run(" decidí salir de nuestro hogar, sin regresar y por ende no reanudar nuestra convivencia común. ")                    
        elif 'ontraparte' in data_load['abandonaHogarComun']:
            parrafo1.add_run(data_load['honorificoContraparte'])
            parrafo1.add_run(" ")                    
            parrafo1.add_run(data_load['nombreContraparte'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoPaternoContraparte'])
            parrafo1.add_run(" ")
            parrafo1.add_run(data_load['apellidoMaternoContraparte'])                
            parrafo1.add_run(" decide salir de nuestro hogar, sin regresar y por ende no reanudar nuestra convivencia común. ")
        parrafo1.add_run("Producto de esta unión ")
 

        if data_load['hijosEntrePartes']==1:
            parrafo1.add_run(str(numeroParrafo))
            if 'F' in data_load['sexoDeHijoEntreParte1']:
                parrafo1.add_run(" nació nuestra hija ")
            elif 'M' in data_load['sexoDeHijoEntreParte1']:
                parrafo1.add_run(" nació nuestro hijo ")
            parrafo1.add_run(data_load['nombreDeHijoEntreParte1'])
            parrafo1.add_run(", RUT ")
            parrafo1.add_run(data_load['rutDeHijoEntreParte1'])
            parrafo1.add_run(", de actuales ")
            parrafo1.add_run(data_load['edadDeHijoEntreParte1'])
            parrafo1.add_run(" años de edad.")
        elif data_load['hijosEntrePartes']>1:                 
            parrafo1.add_run(" nacieron ")
            if data_load['hijosEntrePartes'] == 2:
                if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2']:
                    parrafo1.add_run("nuestras hijas: ")
                else:
                    parrafo1.add_run("nuestros hijos: ")
            elif data_load['hijosEntrePartes'] == 3:
                if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3']:
                    parrafo1.add_run("nuestras hijas: ")
                else:
                    parrafo1.add_run("nuestros hijos: ")
            elif data_load['hijosEntrePartes'] == 4:
                if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3'] and 'F' in data_load['sexoDeHijoEntreParte4']:
                    parrafo1.add_run("nuestras hijas: ")
                else:
                    parrafo1.add_run("nuestros hijos: ")            
            elif data_load['hijosEntrePartes'] == 5:
                if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3'] and 'F' in data_load['sexoDeHijoEntreParte4'] and 'F' in data_load['sexoDeHijoEntreParte5']:
                    parrafo1.add_run("nuestras hijas: ")
                else:
                    parrafo1.add_run("nuestros hijos: ")            
            parrafoHijos=document.add_paragraph()
            parrafoHijos.add_run("- ")
            parrafoHijos.add_run(data_load['nombreDeHijoEntreParte1'])
            parrafoHijos.add_run(", RUT ")
            parrafoHijos.add_run(data_load['rutDeHijoEntreParte1'])
            parrafoHijos.add_run(", de actuales ")
            parrafoHijos.add_run(data_load['edadDeHijoEntreParte1'])
            parrafoHijos.add_run(" años de edad.")
            parrafoHijos=document.add_paragraph()
            parrafoHijos.add_run("- ")
            parrafoHijos.add_run(data_load['nombreDeHijoEntreParte2'])
            parrafoHijos.add_run(", RUT ")
            parrafoHijos.add_run(data_load['rutDeHijoEntreParte2'])
            parrafoHijos.add_run(", de actuales ")
            parrafoHijos.add_run(data_load['edadDeHijoEntreParte2'])
            parrafoHijos.add_run(" años de edad.")
            if 'a' in data_load['nombreDeHijoEntreParte3'] or 'e' in data_load['nombreDeHijoEntreParte3'] or 'i' in data_load['nombreDeHijoEntreParte3'] or 'o' in data_load['nombreDeHijoEntreParte3'] or 'u' in data_load['nombreDeHijoEntreParte3']:
                parrafoHijos=document.add_paragraph()
                parrafoHijos.add_run("- ")
                parrafoHijos.add_run(data_load['nombreDeHijoEntreParte3'])
                parrafoHijos.add_run(", RUT ")
                parrafoHijos.add_run(data_load['rutDeHijoEntreParte3'])
                parrafoHijos.add_run(", de actuales ")
                parrafoHijos.add_run(data_load['edadDeHijoEntreParte3'])
                parrafoHijos.add_run(" años de edad.")
            if 'a' in data_load['nombreDeHijoEntreParte4'] or 'e' in data_load['nombreDeHijoEntreParte4'] or 'i' in data_load['nombreDeHijoEntreParte4'] or 'o' in data_load['nombreDeHijoEntreParte4'] or 'u' in data_load['nombreDeHijoEntreParte4']:
                parrafoHijos=document.add_paragraph()
                parrafoHijos.add_run("- ")
                parrafoHijos.add_run(data_load['nombreDeHijoEntreParte4'])
                parrafoHijos.add_run(", RUT ")
                parrafoHijos.add_run(data_load['rutDeHijoEntreParte4'])
                parrafoHijos.add_run(", de actuales ")
                parrafoHijos.add_run(data_load['edadDeHijoEntreParte4'])
                parrafoHijos.add_run(" años de edad.")
            if 'a' in data_load['nombreDeHijoEntreParte5'] or 'e' in data_load['nombreDeHijoEntreParte5'] or 'i' in data_load['nombreDeHijoEntreParte5'] or 'o' in data_load['nombreDeHijoEntreParte5'] or 'u' in data_load['nombreDeHijoEntreParte5']:
                parrafoHijos=document.add_paragraph()
                parrafoHijos.add_run("- ")
                parrafoHijos.add_run(data_load['nombreDeHijoEntreParte5'])
                parrafoHijos.add_run(", RUT ")
                parrafoHijos.add_run(data_load['rutDeHijoEntreParte5'])
                parrafoHijos.add_run(", de actuales ")
                parrafoHijos.add_run(data_load['edadDeHijoEntreParte5'])
                parrafoHijos.add_run(" años de edad.")     

    parrafo1=document.add_paragraph()
    parrafo1.add_run(" Resulta pertinente señalar que, desde la fecha de separación y hasta")
    parrafo1.add_run(" la interposición de esta demanda, ")
    parrafo1.add_run(data_load['conectorContraparte'])
    parrafo1.add_run(" demandad")
    parrafo1.add_run(data_load['generoContraparte'])
    if data_load['pagoAlimentosContraparte']=='Nada':
        parrafo1.add_run(" no ha realizado ningún aporte para la mantención de nuestr")
        #incluir instrucción si son varios, o hija, o hijas
        parrafo1.add_run(", como tampoco ha participado de su crianza.")
    elif data_load['pagoAlimentosContraparte']=='Ayuda ocasional o errática':
        parrafo1.add_run(" tan solo ha realizado aportes esporádicos para la mantención de nuestr")
        #incluir instrucción si son varios, o hija, o hijas
        parrafo1.add_run(", como tampoco ha participado de su crianza.")    
    elif data_load['pagoAlimentosContraparte']=='Menos de lo fijado':
        parrafo1.add_run(", si bien ha efectuado aportes en dinero para la mantención de nuestr")
        #incluir instrucción si son varios, o hija, o hijas
        parrafo1.add_run(", lo ha hecho en menor medida de lo que le corresponde.")
        parrafo1.add_run(" En adición, no ha participado ni se ha preocupado de su crianza.")       
    
    

       
       
    document.add_paragraph()
    parrafo9=document.add_paragraph()
    parrafo9.add_run("EL DERECHO:").bold=True
    document.add_paragraph()

    numeroParrafo=0

    numeroParrafo+=1

    parrafo10=document.add_paragraph()
    parrafo10.add_run(str(numeroParrafo))
    parrafo10.add_run('El artículo 321 N° 2 del Código Civil dispone que "Se deben alimentos: N°2.- A los descendientes."')

    formatoParrafo10=parrafo10.paragraph_format
    formatoParrafo10.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


    numeroParrafo+=1
    parrafo1=document.add_paragraph()
    parrafo1.add_run(str(numeroParrafo))
    parrafo1.add_run('Por su parte, el artículo 323 inciso 1º del mismo cuerpo legal dispone que “Los alimentos deben habilitar al alimentado para subsistir modestamente de un modo correspondiente a su posición social”.')

    formatoParrafo10=parrafo10.paragraph_format
    formatoParrafo10.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    
    document.add_paragraph()
    parrafo1=document.add_paragraph()
    parrafo1.add_run("POR TANTO:").bold=True
    document.add_paragraph()    
    
    
    numeroParrafo+=1
    parrafo1=document.add_paragraph()
    parrafo1.add_run(str(numeroParrafo))
    parrafo1.add_run('En atención a los hechos expuestos y lo dispuesto en los artículos 321 y siguientes del Código Civil,')
    parrafo1.add_run(' el artículo 3º de la ley Nº 14.908 modificada en lo pertinente, las disposiciones de las leyes Nº 19.968')
    parrafo1.add_run(' y 16.618,  así como los artículos 254 y siguientes del Código de Procedimiento Civil, y demás normas')
    parrafo1.add_run(' pertinentes.')
        
    
    formatoParrafo10=parrafo10.paragraph_format
    formatoParrafo10.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


    parrafo1=document.add_paragraph()
    parrafo1.add_run("SOLICITO A S.S.:").bold=True
    parrafo1.add_run(" Se sirva tener por interpuesta demanda de alimentos menores en contra de ")
    parrafo1.add_run(data_load['honorificoContraparte'])
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['nombreContraparte'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoContraparte'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoContraparte'].upper()).bold=True    
    parrafo1.add_run(", ya individualizad")
    parrafo1.add_run(data_load['generoContraparte'])
    parrafo1.add_run(", acogerla a tramitación y en definitiva dar lugar a ella en todas sus partes condenándol")
    parrafo1.add_run(data_load['generoContraparte'])   
    parrafo1.add_run(" a pagar una pensión alimenticia a favor de nuestr")
#nuestro hijo/ nuestra hija/ nuestros hijos/ nuestras hijas
    parrafo1.add_run(", ascendiente a la suma de ")
    parrafo1.add_run("$").bold=True 
    parrafo1.add_run(data_load['cuantiaAlimentosMenores']).bold=True
    parrafo1.add_run('(').bold=True
    parrafo1.add_run(numeralizar(data_load['cuantiaAlimentosMenores'])).bold=True
    parrafo1.add_run(')').bold=True
    parrafo1.add_run(" pesos mensuales, o lo que SS., estime de justicia, de acuerdo al mérito del procedimiento, la cual se reajustará")
    parrafo1.add_run(" conforme a la variación del ingreso mínimo remuneracional. y, que deberá ser pagada mensualmente")
    parrafo1.add_run(" mediante depósito en la cuenta de ahorro a la vista que se abrirá al efecto.")


    parrafo1=document.add_paragraph()
    parrafo1.add_run("PRIMER OTROSÍ:").bold=True
    parrafo1.add_run(" Atendido lo dispuesto en el artículo 5 de la ley Nº 14.908, sobre Abandono de Familia y pago de pensiones alimenticias,")
    parrafo1.add_run(" así como en el artículo 327 del Código Civil, se cumplen los requerimientos establecidos por el legislador para")
    parrafo1.add_run(" la procedencia de los alimentos provisorios.")


    
    document.add_paragraph("")

    parrafo1=document.add_paragraph()
    parrafo1.add_run("POR TANTO,").bold=True
    parrafo1=document.add_paragraph()
    parrafo1.add_run("SOLICITO A S.S.: ").bold=True
    parrafo1.add_run(" se sirva decretar alimentos provisorios a favor de")
    #el, la, los las, menor o menores
    parrafo1.add_run("de autos por la suma de $")
    parrafo1.add_run(data_load['cuantiaAlimentosMenores']).bold=True 
    parrafo1.add_run("(").bold=True 
    parrafo1.add_run(numeralizar(data_load['cuantiaAlimentosMenores'])).bold=True 
    parrafo1.add_run(")").bold=True 
    parrafo1.add_run(" pesos mensuales").bold=True 
    parrafo1.add_run(", o la suma que S.S. estime en justicia regular.")
    
    
    
    
    #DOCUMENTOS
    parrafo12=document.add_paragraph()
    parrafo12.add_run("SEGUNDO OTROSÍ:").bold=True
    parrafo12.add_run(" Sírvase SS., tener por acompañados con citación los siguientes documentos, para ser introducidos al presente juicio en la audiencia pertinente:")
    document.add_paragraph("")
    #agregar documentos normales, probablemente certificado de nacimiento y lista de gastos.
    parrafo12=document.add_paragraph()
    parrafo12.add_run("    - Certificado de matrimonio entre ")
    parrafo12.add_run(data_load['honorificoPatrocinado'])
    parrafo12.add_run(" ")                    
    parrafo12.add_run(data_load['nombrePatrocinado'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoPaternoPatrocinado'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoMaternoPatrocinado'])
    parrafo12.add_run(" y ")                    
    parrafo12.add_run(data_load['honorificoContraparte'])
    parrafo12.add_run(" ")                    
    parrafo12.add_run(data_load['nombreContraparte'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoPaternoContraparte'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoMaternoContraparte'])
    parrafo12.add_run(".")                    
                
    if 'a' in data_load['nombreDeHijoEntreParte1'] or 'e' in data_load['nombreDeHijoEntreParte1'] or 'i' in data_load['nombreDeHijoEntreParte1'] or 'o' in data_load['nombreDeHijoEntreParte1'] or 'u' in data_load['nombreDeHijoEntreParte1']:
        parrafo12=document.add_paragraph()                    
        parrafo12.add_run("    - Certificado de nacimiento de ")
        parrafo12.add_run(data_load['nombreDeHijoEntreParte1'])
        parrafo12.add_run(".")                            
    if 'a' in data_load['nombreDeHijoEntreParte2'] or 'e' in data_load['nombreDeHijoEntreParte2'] or 'i' in data_load['nombreDeHijoEntreParte2'] or 'o' in data_load['nombreDeHijoEntreParte2'] or 'u' in data_load['nombreDeHijoEntreParte2']:
        parrafo12=document.add_paragraph()                                        
        parrafo12.add_run("    - Certificado de nacimiento de ")
        parrafo12.add_run(data_load['nombreDeHijoEntreParte2'])
        parrafo12.add_run(".")                       
    if 'a' in data_load['nombreDeHijoEntreParte3'] or 'e' in data_load['nombreDeHijoEntreParte3'] or 'i' in data_load['nombreDeHijoEntreParte3'] or 'o' in data_load['nombreDeHijoEntreParte3'] or 'u' in data_load['nombreDeHijoEntreParte3']:
        parrafo12=document.add_paragraph()                                        
        parrafo12.add_run("    - Certificado de nacimiento de ")
        parrafo12.add_run(data_load['nombreDeHijoEntreParte3'])
        parrafo12.add_run(".")                     
    if 'a' in data_load['nombreDeHijoEntreParte4'] or 'e' in data_load['nombreDeHijoEntreParte4'] or 'i' in data_load['nombreDeHijoEntreParte4'] or 'o' in data_load['nombreDeHijoEntreParte4'] or 'u' in data_load['nombreDeHijoEntreParte4']:
        parrafo12=document.add_paragraph()                                        
        parrafo12.add_run("    - Certificado de nacimiento de ")
        parrafo12.add_run(data_load['nombreDeHijoEntreParte4'])
        parrafo12.add_run(".")                     
    if 'a' in data_load['nombreDeHijoEntreParte5'] or 'e' in data_load['nombreDeHijoEntreParte5'] or 'i' in data_load['nombreDeHijoEntreParte5'] or 'o' in data_load['nombreDeHijoEntreParte5'] or 'u' in data_load['nombreDeHijoEntreParte5']:
        parrafo12=document.add_paragraph()                                        
        parrafo12.add_run("    - Certificado de nacimiento de ")
        parrafo12.add_run(data_load['nombreDeHijoEntreParte5'])
        parrafo12.add_run(".")                     



                     
                          
    parrafo1=document.add_paragraph()
    parrafo1.add_run("POR TANTO,").bold=True
    parrafo1=document.add_paragraph()
    parrafo1.add_run("SOLICITO A S.S.: ").bold=True
    parrafo1.add_run("se sirva tenerlos por acompañados.")

    parrafo1=document.add_paragraph()
    parrafo1.add_run("TERCER OTROSÍ:").bold=True
    parrafo1.add_run(" Ruego a U.S., autorizarme para litigar en forma electrónica y ser notificad")
    parrafo1.add_run(data_load['generoPatrocinado'])
    parrafo1.add_run(" de las resoluciones que se dicten en la presente causa a través del correo electrónico ")
    parrafo1.add_run(data_load['emailAbogadoPatrocinado'])
    parrafo1.add_run(".")
    
    parrafo1=document.add_paragraph()
    parrafo1.add_run("CUARTO OTROSÍ:").bold=True
    parrafo1.add_run(" Pido a U.S., tener presente que designo abogad")
    parrafo1.add_run(data_load['generoAbogadoPatrocinado'])
    parrafo1.add_run(" patrocinante y confiero poder a ")
    parrafo1.add_run(data_load['honorificoAbogadoPatrocinado'])
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['nombreAbogadoPatrocinado'].upper())
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoAbogadoPatrocinado'].upper())
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoAbogadoPatrocinado'].upper())                   
    parrafo1.add_run(", cédula de identidad Nº ")
    parrafo1.add_run(data_load['rutAbogadoPatrocinado'])
    parrafo1.add_run(", abogad")
    parrafo1.add_run(data_load['generoAbogadoPatrocinado'])
    parrafo1.add_run(" habilitad")
    parrafo1.add_run(data_load['generoAbogadoPatrocinado'])
    parrafo1.add_run(" para el ejercicio de la profesion, con domicilio en ")
    parrafo1.add_run(data_load['domicilioAbogadoPatrocinado'])
    parrafo1.add_run(" número ")
    parrafo1.add_run(data_load['numeroDomicilioAbogadoPatrocinado'])     
    parrafo1.add_run(" comuna de ")
    parrafo1.add_run(data_load['comunaAbogadoPatrocinado'])    
    if data_load['facultadesAbogadoPatrocinado'] == 'Sí':
        parrafo1.add_run(' para que ')
        parrafo1.add_run(data_load['conectorPatrocinado'])
        parrafo1.add_run(' represente con las facultades de ambos incisos del artículo 7° del Código de Procedimiento Civil, en especial, las de percibir, avenir y transigir,')    
    parrafo12.add_run(" que firma en señal de aceptación")
    if 'a' in data_load['nombreApoderadoPatrocinado'] or 'e' in data_load['nombreApoderadoPatrocinado'] or 'i' in data_load['nombreApoderadoPatrocinado'] or 'o' in data_load['nombreApoderadoPatrocinado'] or 'u' in data_load['nombreApoderadoPatrocinado']:
        parrafo1.add_run("; y también ")
        parrafo1.add_run(conector2ApoderadoPatrocinado)
        parrafo1.add_run(" ")
        parrafo1.add_run(profesionApoderadoPatrocinado)
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['nombreApoderadoPatrocinado'].upper())
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoPaternoApoderadoPatrocinado'].upper())
        parrafo1.add_run(" ")
        parrafo1.add_run(data_load['apellidoMaternoApoderadoPatrocinado'].upper())                                   
        parrafo1.add_run(", cédula de identidad Nº ")
        parrafo1.add_run(data_load['rutApoderadoPatrocinado'])
        parrafo1.add_run(", del mismo domicilio y forma de notificación")
        if data_load['facultadesApoderadoPatrocinado'] == 'Sí':
            parrafo1.add_run(', para que me')
            parrafo1.add_run(' represente con las facultades de ambos incisos del artículo 7° del Código de Procedimiento Civil, en especial, las de percibir, avenir y transigir,')
        else:
            parrafo1.add_run('.')        
    else:
        parrafo1.add_run(".")    



    document.add_paragraph("")

    parrafo1=document.add_paragraph()
    parrafo1.add_run("POR TANTO,").bold=True

    parrafo1=document.add_paragraph()
    parrafo1.add_run("SOLICITO A S.S.: ").bold=True
    parrafo1.add_run("se sirva tener presente patrocinio y poder conferido.")



    formatoParrafo1=parrafo1.paragraph_format
    formatoParrafo1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save('alMen'+data_load['apellidoPaternoPatrocinado']+data_load['apellidoPaternoContraparte']+'.docx')

    print("I am a zombie, ZOMBIE!!")
