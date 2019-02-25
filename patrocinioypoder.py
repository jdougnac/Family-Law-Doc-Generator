from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Length
from docx.text.run import Font
from docx.table import Table
from tkinter import *
from tkinter.filedialog import *
import json


def patrocinioYPoder():
    with open(filedialog.askopenfilename()) as fp:
            data_load=json.load(fp)

      
#DOCUMENTO PROPIAMENTE TAL

    document=Document("documentoBase.docx")
    Font.name = 'Century Gothic'
    Font.size = Pt(16)
    
    suma="patrocinio y poder"
    encabezado1=document.add_paragraph()
    encabezado1.add_run(suma.upper()).bold=True


    encabezado2=document.add_paragraph()
    encabezado2.add_run("S.J. DE FAMILIA DE SANTIAGO (").bold=True
    encabezado2.add_run(data_load['juzgado'].upper()).bold=True
    encabezado2.add_run(")")
    formatoEncabezado2=encabezado2.paragraph_format
    formatoEncabezado2.alignment=WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()

    parrafo1=document.add_paragraph()
    parrafo1.add_run("     ")
    parrafo1.add_run(data_load['nombrePatrocinado'].upper())
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoPatrocinado'].upper())
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoPatrocinado'].upper())
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['nacionalidadPatrocinado'])
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['estadoCivilPatrocinado'])
    parrafo1.add_run(", cedula nacional de identidad numero ")
    parrafo1.add_run(data_load['rutPatrocinado'])
    parrafo1.add_run(", con domicilio en ")
    parrafo1.add_run(data_load['domicilioPatrocinado'])
    parrafo1.add_run(" número ")
    parrafo1.add_run(data_load['numeroDomicilioPatrocinado'])     
    parrafo1.add_run(", comuna de ")
    parrafo1.add_run(data_load['comunaPatrocinado'])
    parrafo1.add_run(", parte ")
    parrafo1.add_run(data_load['situacionProcesalPatrocinado'])
    parrafo1.add_run(" en autos sobre ")
    parrafo1.add_run(data_load['materiaCausa'])
    parrafo1.add_run(", caratulados ")
    parrafo1.add_run(data_load['caratulaCausa'])
    parrafo1.add_run(", en causa RIT ")
    parrafo1.add_run(data_load['ritDeCausa'])
    parrafo1.add_run(", a U.S., respetuosamente digo:")

    formatoParrafo1=parrafo1.paragraph_format
    formatoParrafo1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    formatoParrafo1.line_spacing=1.5


    document.add_paragraph()

    parrafo2=document.add_paragraph()
    parrafo2.add_run("     ")
    parrafo2.add_run("Que en este acto, confiero patrocinio y poder ")
    parrafo2.add_run(data_load['conector2AbogadoPatrocinado'])
    parrafo2.add_run(" ")
    parrafo2.add_run(data_load['profesionAbogadoPatrocinado'])
    parrafo2.add_run(" ")
    parrafo2.add_run(data_load['honorificoAbogadoPatrocinado'])
    parrafo2.add_run(" ")    
    parrafo2.add_run(data_load['nombreAbogadoPatrocinado'].upper())
    parrafo2.add_run(" ")
    parrafo2.add_run(data_load['apellidoPaternoAbogadoPatrocinado'].upper())
    parrafo2.add_run(" ")
    parrafo2.add_run(data_load['apellidoMaternoAbogadoPatrocinado'].upper())
    parrafo2.add_run(", cedula nacional de identidad numero ")
    parrafo2.add_run(data_load['rutAbogadoPatrocinado'])
    parrafo2.add_run(", facultad")
    parrafo2.add_run(data_load['generoAbogadoPatrocinado'])
    parrafo2.add_run(" para el ejercicio de la profesion, quien fija su domicilio en ")
    parrafo2.add_run(data_load['domicilioAbogadoPatrocinado'])
    parrafo2.add_run(" número ")
    parrafo2.add_run(data_load['numeroDomicilioAbogadoPatrocinado'])     
    parrafo2.add_run(", comuna de ")
    parrafo2.add_run(data_load['comunaAbogadoPatrocinado'])
    parrafo2.add_run(", y quien firma junto a mi en señal de aceptacion, solicitando desde ya sean incorporados al sistema SITFA, señalando como forma de notificacion el correo electronico: ")
    parrafo2.add_run(data_load['emailAbogadoPatrocinado'])
    parrafo2.add_run(".")

    formatoParrafo2=parrafo2.paragraph_format
    formatoParrafo2.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    formatoParrafo2.line_spacing=1.5

    parrafo3=document.add_paragraph("POR TANTO,").bold=True
    parrafo5=document.add_paragraph("PIDO A US.,").bold=True
    parrafo6=document.add_paragraph("tenerlo presente")


#cambiar el save del documento

    document.save('PatPoder'+data_load['apellidoPaternoPatrocinado']+data_load['apellidoPaternoContraparte']+'.docx')
    
    print("I am a zombie, ZOMBIE!!")
