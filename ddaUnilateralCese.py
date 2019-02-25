from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Length
from docx.text.run import Font
from docx.table import Table
from tkinter import *
from tkinter.filedialog import *
import json


def divorcioUnilateralCese():
    with open(filedialog.askopenfilename()) as fp:
            data_load=json.load(fp)


        
    #DOCUMENTO PROPIAMENTE TAL

    document=Document("documentoBase.docx")
    Font.name = 'Century Gothic'
    Font.size = Pt(16)



    tablaDatos=document.add_table(rows=3, cols=2)

    proc1=tablaDatos.cell(0,0)
    proc1.paragraphs[0].add_run('PROCEDIMIENTO').bold = True

    proc11=tablaDatos.cell(0,1)
    proc11.text=': ORDINARIO'


    proc2=tablaDatos.cell(1,0)
    proc2.paragraphs[0].add_run('MATERIA').bold = True
    proc22=tablaDatos.cell(1,1)
    proc22.text=': DIVORCIO DE MUTUO ACUERDO'

    proc3=tablaDatos.cell(2,0)
    proc3.paragraphs[0].add_run('FECHA DE MATRIMONIO').bold = True

    proc33=tablaDatos.cell(2,1)
    proc33.paragraphs[0].add_run(': ')    
    proc33.paragraphs[0].add_run(data_load['diaDeMatrimonio'])
    proc33.paragraphs[0].add_run(' de ')
    proc33.paragraphs[0].add_run(data_load['mesDeMatrimonio'])
    proc33.paragraphs[0].add_run(' de ')
    proc33.paragraphs[0].add_run(data_load['añoDeMatrimonio'])
        
    document.add_paragraph('') 

    tablaPatrocinado=document.add_table(rows=3, cols=2)

    proc1=tablaPatrocinado.cell(0,0)
    proc1.paragraphs[0].add_run('DEMANDANTE').bold = True
    proc11=tablaPatrocinado.cell(0,1)
    proc11.paragraphs[0].add_run(': ')
    proc11.paragraphs[0].add_run(data_load['nombrePatrocinado'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoPaternoPatrocinado'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoMaternoPatrocinado'])      

    proc2=tablaPatrocinado.cell(1,0)
    proc2.paragraphs[0].add_run('RUT').bold = True
    proc22=tablaPatrocinado.cell(1,1)
    proc22.paragraphs[0].add_run(': ')
    proc22.paragraphs[0].add_run(data_load['rutPatrocinado'])

    proc3=tablaPatrocinado.cell(2,0)
    proc3.paragraphs[0].add_run('DOMICILIO').bold = True

    proc33=tablaPatrocinado.cell(2,1)
    proc33.paragraphs[0].add_run(': ')    
    proc33.paragraphs[0].add_run(data_load['domicilioPatrocinado'])
    proc33.paragraphs[0].add_run(" ")
    proc33.paragraphs[0].add_run(data_load['numeroDomicilioPatrocinado']) 
    proc33.paragraphs[0].add_run(", comuna de ")
    proc33.paragraphs[0].add_run(data_load['comunaPatrocinado'])    
    document.add_paragraph('')   

    tablaAbogadoPatrocinado=document.add_table(rows=4, cols=2)

    proc1=tablaAbogadoPatrocinado.cell(0,0)
    proc1.paragraphs[0].add_run(data_load['profesionAbogadoPatrocinado'].upper()).bold = True
    proc1.paragraphs[0].add_run(' PATROCINANTE').bold = True
    proc11=tablaAbogadoPatrocinado.cell(0,1)
    proc11.paragraphs[0].add_run(': ')
    proc11.paragraphs[0].add_run(data_load['nombreAbogadoPatrocinado'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoPaternoAbogadoPatrocinado'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoMaternoAbogadoPatrocinado'])      

    proc2=tablaAbogadoPatrocinado.cell(1,0)
    proc2.paragraphs[0].add_run('RUT').bold = True
    proc22=tablaAbogadoPatrocinado.cell(1,1)
    proc22.paragraphs[0].add_run(': ')
    proc22.paragraphs[0].add_run(data_load['rutAbogadoPatrocinado'])

    proc3=tablaAbogadoPatrocinado.cell(2,0)
    proc3.paragraphs[0].add_run('DOMICILIO').bold = True

    proc33=tablaAbogadoPatrocinado.cell(2,1)
    proc33.paragraphs[0].add_run(': ')    
    proc33.paragraphs[0].add_run(data_load['domicilioAbogadoPatrocinado'])
    proc33.paragraphs[0].add_run(" ")    
    proc33.paragraphs[0].add_run(data_load['numeroDomicilioAbogadoPatrocinado']) 
    proc33.paragraphs[0].add_run(", comuna de ")
    proc33.paragraphs[0].add_run(data_load['comunaAbogadoPatrocinado'])       

    proc4=tablaAbogadoPatrocinado.cell(3,0)
    proc4.paragraphs[0].add_run('CORREO ELECTRÓNICO').bold = True

    proc44=tablaAbogadoPatrocinado.cell(3,1)
    proc44.paragraphs[0].add_run(': ')    
    proc44.paragraphs[0].add_run(data_load['emailAbogadoPatrocinado'])

    document.add_paragraph('')    


    if 'a' in data_load['nombreApoderadoPatrocinado'] or 'e' in data_load['nombreApoderadoPatrocinado'] or 'i' in data_load['nombreApoderadoPatrocinado'] or 'o' in data_load['nombreApoderadoPatrocinado'] or 'u' in data_load['nombreApoderadoPatrocinado']:
        tablaApoderadoPatrocinado=document.add_table(rows=4, cols=2)

        proc1=tablaApoderadoPatrocinado.cell(0,0)
        proc1.paragraphs[0].add_run(profesionApoderadoPatrocinado.upper()).bold = True
        proc1.paragraphs[0].add_run(' APODERAD').bold = True
        proc1.paragraphs[0].add_run(data_load['generoApoderadoPatrocinado'].upper()).bold = True        
        proc11=tablaApoderadoPatrocinado.cell(0,1)
        proc11.paragraphs[0].add_run(': ')
        proc11.paragraphs[0].add_run(data_load['nombreApoderadoPatrocinado'])
        proc11.paragraphs[0].add_run(' ')
        proc11.paragraphs[0].add_run(data_load['apellidoPaternoApoderadoPatrocinado'])
        proc11.paragraphs[0].add_run(' ')
        proc11.paragraphs[0].add_run(data_load['apellidoMaternoApoderadoPatrocinado'])      

        proc2=tablaApoderadoPatrocinado.cell(1,0)
        proc2.paragraphs[0].add_run('RUT').bold = True
        proc22=tablaApoderadoPatrocinado.cell(1,1)
        proc22.paragraphs[0].add_run(': ')
        proc22.paragraphs[0].add_run(data_load['rutApoderadoPatrocinado'])

        proc3=tablaApoderadoPatrocinado.cell(2,0)
        proc3.paragraphs[0].add_run('DOMICILIO').bold = True

        proc33=tablaApoderadoPatrocinado.cell(2,1)
        proc33.paragraphs[0].add_run(': ')    
        proc33.paragraphs[0].add_run(data_load['domicilioApoderadoPatrocinado'])
        proc33.paragraphs[0].add_run(" ")    
        proc33.paragraphs[0].add_run(data_load['numeroDomicilioApoderadoPatrocinado']) 
        proc33.paragraphs[0].add_run(", comuna de ")
        proc33.paragraphs[0].add_run(data_load['comunaApoderadoPatrocinado'])          

        proc4=tablaApoderadoPatrocinado.cell(3,0)
        proc4.paragraphs[0].add_run('CORREO ELECTRÓNICO').bold = True

        proc44=tablaApoderadoPatrocinado.cell(3,1)
        proc44.paragraphs[0].add_run(': ')    
        proc44.paragraphs[0].add_run(data_load['emailApoderadoPatrocinado'])      


    document.add_paragraph('')   

    tablaContraparte=document.add_table(rows=3, cols=2)

    proc1=tablaContraparte.cell(0,0)
    proc1.paragraphs[0].add_run('DEMANDAD').bold = True
    proc1.paragraphs[0].add_run(data_load['generoContraparte'].upper()).bold = True
    proc11=tablaContraparte.cell(0,1)
    proc11.paragraphs[0].add_run(': ')
    proc11.paragraphs[0].add_run(data_load['nombreContraparte'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoPaternoContraparte'])
    proc11.paragraphs[0].add_run(' ')
    proc11.paragraphs[0].add_run(data_load['apellidoMaternoContraparte'])      

    proc2=tablaContraparte.cell(1,0)
    proc2.paragraphs[0].add_run('RUT').bold = True
    proc22=tablaContraparte.cell(1,1)
    proc22.paragraphs[0].add_run(': ')
    proc22.paragraphs[0].add_run(data_load['rutContraparte'])

    proc3=tablaContraparte.cell(2,0)
    proc3.paragraphs[0].add_run('DOMICILIO').bold = True

    proc33=tablaContraparte.cell(2,1)
    proc33.paragraphs[0].add_run(': ')    
    proc33.paragraphs[0].add_run(data_load['domicilioContraparte'])
    proc33.paragraphs[0].add_run(" ")    
    proc33.paragraphs[0].add_run(data_load['numeroDomicilioContraparte'])
    proc33.paragraphs[0].add_run(", comuna de ")
    proc33.paragraphs[0].add_run(data_load['comunaContraparte'])      

    document.add_paragraph('')   

    
    resumen1=document.add_paragraph()
    resumen1.add_run("EN LO PRINCIPAL: ").bold=True
    resumen1.add_run("Demanda de divorcio unilateral; ")
    resumen1.add_run("PRIMER OTROSÍ: ").bold=True
    resumen1.add_run("Acompaña documentos; ")
    resumen1.add_run("SEGUNDO OTROSÍ: ").bold=True
    resumen1.add_run("Medio de notificación que indica; ")
    resumen1.add_run("TERCER OTROSÍ: ").bold=True
    resumen1.add_run("Patrocinio y poder; ")


    document.add_paragraph("")

    encabezado2=document.add_paragraph()
    encabezado2.add_run("S.J. DE FAMILIA DE SANTIAGO (").bold=True
    encabezado2.add_run(data_load['juzgado']).bold=True
    encabezado2.add_run(")").bold=True
    formatoEncabezado2=encabezado2.paragraph_format
    formatoEncabezado2.alignment=WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")


    parrafo1=document.add_paragraph()

    parrafo1.add_run("    ")
    parrafo1.add_run(data_load['nombrePatrocinado'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoPatrocinado'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoPatrocinado'].upper()).bold=True
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['estadoCivilPatrocinado'])
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['nacionalidadPatrocinado'])
    parrafo1.add_run(", cédula de identidad Nº ")
    parrafo1.add_run(data_load['rutPatrocinado'])
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['profesionPatrocinado'])
    parrafo1.add_run(", domiciliad")
    parrafo1.add_run(data_load['generoPatrocinado'])
    parrafo1.add_run(" en ")
    parrafo1.add_run(data_load['domicilioPatrocinado'])
    parrafo1.add_run(" número ")
    parrafo1.add_run(data_load['numeroDomicilioPatrocinado']) 
    parrafo1.add_run(", comuna de ")
    parrafo1.add_run(data_load['comunaPatrocinado'])
    parrafo1.add_run("; a U.S., respetuosamente digo:")

    formatoParrafo1=parrafo1.paragraph_format
    formatoParrafo1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    formatoParrafo1.line_spacing=1.5


    parrafo2=document.add_paragraph()
    parrafo2.add_run("      ")
    parrafo2.add_run("Que, por este acto interpongo demanda de Divorcio Unilateral en contra de mi cónyuge ")
    parrafo2.add_run(data_load['honorificoContraparte'])
    parrafo2.add_run(" ")    
    parrafo2.add_run(data_load['nombreContraparte'].upper()).bold=True
    parrafo2.add_run(" ")
    parrafo2.add_run(data_load['apellidoPaternoContraparte'].upper()).bold=True
    parrafo2.add_run(" ")
    parrafo2.add_run(data_load['apellidoMaternoContraparte'].upper()).bold=True
    parrafo2.add_run(", ")
    parrafo2.add_run(data_load['estadoCivilContraparte'])
    parrafo2.add_run(", ")
    parrafo2.add_run(data_load['nacionalidadContraparte'])
    parrafo2.add_run(", cédula de identidad Nº ")
    parrafo2.add_run(data_load['rutContraparte'])
    parrafo2.add_run(", ")
    parrafo2.add_run(data_load['profesionContraparte'])
    parrafo2.add_run(", domiciliad")
    parrafo2.add_run(data_load['generoContraparte'])
    parrafo2.add_run(" en ")
    parrafo2.add_run(data_load['domicilioContraparte'])
    parrafo2.add_run(" número ")
    parrafo2.add_run(data_load['numeroDomicilioContraparte']) 
    parrafo2.add_run(", comuna de ")
    parrafo2.add_run(data_load['comunaContraparte'])
    parrafo2.add_run(", en virtud de los argumentos de hecho y de Derecho que paso a exponer: ")

    formatoParrafo2=parrafo2.paragraph_format
    formatoParrafo2.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    formatoParrafo2.line_spacing=1.5


    document.add_paragraph()
    hechos="LOS HECHOS:"
    parrafo3=document.add_paragraph()
    parrafo3.add_run(hechos).bold=True
    document.add_paragraph()

    numeroParrafo=0
    numeroParrafo+=1
    parrafo3=document.add_paragraph()
    parrafo3.add_run(str(numeroParrafo))
    parrafo3.add_run(".- Con fecha ")
    parrafo3.add_run(data_load['diaDeMatrimonio'])
    parrafo3.add_run(" de ")
    parrafo3.add_run(data_load['mesDeMatrimonio'])
    parrafo3.add_run(" de ")
    parrafo3.add_run(data_load['añoDeMatrimonio'])
    parrafo3.add_run(" contraje matrimonio con ")
    parrafo3.add_run(data_load['honorificoContraparte'])
    parrafo3.add_run(" ")    
    parrafo3.add_run(data_load['nombreContraparte'])
    parrafo3.add_run(" ")
    parrafo3.add_run(data_load['apellidoPaternoContraparte'])
    parrafo3.add_run(" ")
    parrafo3.add_run(data_load['apellidoMaternoContraparte'])
    parrafo3.add_run(" ante Oficial de Registro Civil de la circunscripcion de ")
    parrafo3.add_run(data_load['comunaDeMatrimonio'])
    parrafo3.add_run(", el que fue anotado bajo la inscripción Nº ")
    parrafo3.add_run(data_load['numeroInscripcionMatrimonio'])
    parrafo3.add_run(" de dicho año pactando el régimen matrimonial de ")
    parrafo3.add_run(data_load['regimenMatrimonial'].lower())
    parrafo3.add_run(", el cual no ha sido modificado hasta el día de hoy.")

    formatoParrafo3=parrafo3.paragraph_format
    formatoParrafo3.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    numeroParrafo+=1
    parrafo4=document.add_paragraph()

    if data_load['hijosEntrePartes']==0:
        parrafo4.add_run(str(numeroParrafo))
        parrafo4.add_run(".- Que de nuestra unión no nacieron hijos.")
    elif data_load['hijosEntrePartes']==1:
        parrafo4.add_run(str(numeroParrafo))        
        if 'F' in data_load['sexoDeHijoEntreParte1']:
            parrafo4.add_run("2.- Que de nuestra union nació nuestra hija ")
        elif 'M' in data_load['sexoDeHijoEntreParte1']:
            parrafo4.add_run("Que de nuestra unión nació nuestro hijo ")
        parrafo4.add_run(data_load['nombreDeHijoEntreParte1'])
        parrafo4.add_run(", RUT ")
        parrafo4.add_run(data_load['rutDeHijoEntreParte1'])
        parrafo4.add_run(", de actuales ")
        parrafo4.add_run(data_load['edadDeHijoEntreParte1'])
        parrafo4.add_run(" años de edad.")
    elif data_load['hijosEntrePartes']>1:
        parrafo4.add_run(str(numeroParrafo))                    
        parrafo4.add_run(".- Que de nuestra unión nacieron ")
        if data_load['hijosEntrePartes'] == 2:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2']:
                parrafo4.add_run("nuestras hijas:")
            else:
                parrafo4.add_run("nuestros hijos:")
        elif data_load['hijosEntrePartes'] == 3:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3']:
                parrafo4.add_run("nuestras hijas:")
            else:
                parrafo4.add_run("nuestros hijos:")
        elif data_load['hijosEntrePartes'] == 4:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3'] and 'F' in data_load['sexoDeHijoEntreParte4']:
                parrafo4.add_run("nuestras hijas:")
            else:
                parrafo4.add_run("nuestros hijos:")            
        elif data_load['hijosEntrePartes'] == 5:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3'] and 'F' in data_load['sexoDeHijoEntreParte4'] and 'F' in data_load['sexoDeHijoEntreParte5']:
                parrafo4.add_run("nuestras hijas:")
            else:
                parrafo4.add_run("nuestros hijos:")         
        parrafo3=document.add_paragraph()
        parrafo3.add_run("- ")
        parrafo3.add_run(data_load['nombreDeHijoEntreParte1'])
        parrafo3.add_run(", RUT ")
        parrafo3.add_run(data_load['rutDeHijoEntreParte1'])
        parrafo3.add_run(", de actuales ")
        parrafo3.add_run(data_load['edadDeHijoEntreParte1'])
        parrafo3.add_run(" años de edad.")
        parrafo3=document.add_paragraph()
        parrafo3.add_run("- ")
        parrafo3.add_run(data_load['nombreDeHijoEntreParte2'])
        parrafo3.add_run(", RUT ")
        parrafo3.add_run(data_load['rutDeHijoEntreParte2'])
        parrafo3.add_run(", de actuales ")
        parrafo3.add_run(data_load['edadDeHijoEntreParte2'])
        parrafo3.add_run(" años de edad.")
        if 'a' in data_load['nombreDeHijoEntreParte3'] or 'e' in data_load['nombreDeHijoEntreParte3'] or 'i' in data_load['nombreDeHijoEntreParte3'] or 'o' in data_load['nombreDeHijoEntreParte3'] or 'u' in data_load['nombreDeHijoEntreParte3']:
            parrafo3=document.add_paragraph()
            parrafo3.add_run("- ")
            parrafo3.add_run(data_load['nombreDeHijoEntreParte3'])
            parrafo3.add_run(", RUT ")
            parrafo3.add_run(data_load['rutDeHijoEntreParte3'])
            parrafo3.add_run(", de actuales ")
            parrafo3.add_run(data_load['edadDeHijoEntreParte3'])
            parrafo3.add_run(" años de edad.")
        if 'a' in data_load['nombreDeHijoEntreParte4'] or 'e' in data_load['nombreDeHijoEntreParte4'] or 'i' in data_load['nombreDeHijoEntreParte4'] or 'o' in data_load['nombreDeHijoEntreParte4'] or 'u' in data_load['nombreDeHijoEntreParte4']:
            parrafo3=document.add_paragraph()
            parrafo3.add_run("- ")
            parrafo3.add_run(data_load['nombreDeHijoEntreParte4'])
            parrafo3.add_run(", RUT ")
            parrafo3.add_run(data_load['rutDeHijoEntreParte4'])
            parrafo3.add_run(", de actuales ")
            parrafo3.add_run(data_load['edadDeHijoEntreParte4'])
            parrafo3.add_run(" años de edad.")
        if 'a' in data_load['nombreDeHijoEntreParte5'] or 'e' in data_load['nombreDeHijoEntreParte5'] or 'i' in data_load['nombreDeHijoEntreParte5'] or 'o' in data_load['nombreDeHijoEntreParte5'] or 'u' in data_load['nombreDeHijoEntreParte5']:
            parrafo3=document.add_paragraph()
            parrafo3.add_run("- ")
            parrafo3.add_run(data_load['nombreDeHijoEntreParte5'])
            parrafo3.add_run(", RUT ")
            parrafo3.add_run(data_load['rutDeHijoEntreParte5'])
            parrafo3.add_run(", de actuales ")
            parrafo3.add_run(data_load['edadDeHijoEntreParte5'])
            parrafo3.add_run(" años de edad.")                    
    formatoParrafo4=parrafo4.paragraph_format
    formatoParrafo4.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


    numeroParrafo+=1
    parrafo5=document.add_paragraph()
    parrafo5.add_run(str(numeroParrafo))
    parrafo5.add_run(".- Nuestra convivencia fue normal aproximadamente por ")
    parrafo5.add_run(str(data_load['duracionMatrimonio']))
    parrafo5.add_run(" años, pero al transcurrir el tiempo empezamos a tener desavenencias e incompatibilidades irreconciliables, lo que hizo que nuestra convivencia fuera insoportable para seguir manteniéndola,  razón por la cual nos separamos de hecho el año ")
    parrafo5.add_run(data_load['añoDeCeseConvivencia'])
    parrafo5.add_run(", fecha en que ")
    if 'atrocinado' in data_load['abandonaHogarComun']:               
        parrafo5.add_run(" decidi salir de nuestro hogar, sin regresar y por ende no reanudar nuestra convivencia común. ")                    
    elif 'ontraparte' in data_load['abandonaHogarComun']:
        parrafo5.add_run(data_load['honorificoContraparte'])
        parrafo5.add_run(" ")                    
        parrafo5.add_run(data_load['nombreContraparte'])
        parrafo5.add_run(" ")
        parrafo5.add_run(data_load['apellidoPaternoContraparte'])
        parrafo5.add_run(" ")
        parrafo5.add_run(data_load['apellidoMaternoContraparte'])                
        parrafo5.add_run(" decide salir de nuestro hogar, sin regresar y por ende no reanudar nuestra convivencia común. ")

    formatoParrafo5=parrafo5.paragraph_format
    formatoParrafo5.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY



    if data_load['nuevaRelacionPatrocinado']=='Si' and data_load['nuevaRelacionContraparte']=='Si':
        numeroParrafo+=1
        parrafo6=document.add_paragraph()
        parrafo6.add_run(str(numeroParrafo))
        parrafo6.add_run(".- Con posterioridad a nuestra separación tanto ")
        parrafo6.add_run(data_load['honorificoPatrocinado'])
        parrafo6.add_run(" ")                    
        parrafo6.add_run(data_load['nombrePatrocinado'])
        parrafo6.add_run(" ")
        parrafo6.add_run(data_load['apellidoPaternoPatrocinado'])
        parrafo6.add_run(" ")
        parrafo6.add_run(data_load['apellidoMaternoPatrocinado']) 
        parrafo6.add_run(" como ")
        parrafo6.add_run(data_load['honorificoContraparte'])
        parrafo6.add_run(" ")                    
        parrafo6.add_run(data_load['nombreContraparte'])
        parrafo6.add_run(" ")
        parrafo6.add_run(data_load['apellidoPaternoContraparte'])
        parrafo6.add_run(" ")
        parrafo6.add_run(data_load['apellidoMaternoContraparte']) 
        parrafo6.add_run(" hemos iniciado nuevas relaciones, las cuales se extienden por más de ")
        parrafo6.add_run(data_load['añosRelacionPatrocinado'])
        parrafo6.add_run(" y ")
        parrafo6.add_run(data_load['añosRelacionContraparte'])
        parrafo6.add_run(" años respectivamente. ")
        if data_load['hijosPatrocinado']>0 or data_load['hijosContraparte']>0:
            parrafo6.add_run("De esta relación ")
            if data_load['hijosPatrocinado']==0:
                parrafo6.add_run("no han nacido hijos de parte de ")
            elif data_load['hijosPatrocinado']==1:
                parrafo6.add_run("ha nacido de parte de ")
            elif data_load['hijosPatrocinado']>1:
                parrafo6.add_run("han nacido de parte de ")
            parrafo6.add_run(data_load['honorificoPatrocinado'])
            parrafo6.add_run(" ")
            parrafo6.add_run(data_load['nombrePatrocinado'])
            parrafo6.add_run(" ")
            parrafo6.add_run(data_load['apellidoPaternoPatrocinado'])
            parrafo6.add_run(" ")
            parrafo6.add_run(data_load['apellidoMaternoPatrocinado'])
        if data_load['hijosPatrocinado']==1:
            if 'F' in data_load['sexoDeHijoPatrocinado1']:
                parrafo6.add_run(" su hija ")
            elif 'M' in data_load['sexoDeHijoPatrocinado1']:
                parrafo6.add_run(" su hijo ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado1'])
            parrafo6.add_run(".")

        elif data_load['hijosPatrocinado']==2:
            parrafo6.add_run(str(numeroParrafo))                    
            if 'F' in data_load['sexoDeHijoPatrocinado1'] and 'F' in data_load['sexoDeHijoPatrocinado2']:
                parrafo6.add_run(" sus hijas ")
            else:
                parrafo6.add_run(" sus hijos ")            
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado1'])
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado2'])
            parrafo6.add_run(".")

        elif data_load['hijosPatrocinado']==3:
            parrafo6.add_run(str(numeroParrafo))                    
            if 'F' in data_load['sexoDeHijoPatrocinado1'] and 'F' in data_load['sexoDeHijoPatrocinado2'] and 'F' in data_load['sexoDeHijoPatrocinado3']:
                parrafo6.add_run(" sus hijas ")
            else:
                parrafo6.add_run(" sus hijos ")            
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado1'])
            parrafo6.add_run(", ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado2'])            
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado3'])
            parrafo6.add_run(".")

        elif data_load['hijosPatrocinado']==4:
            parrafo6.add_run(str(numeroParrafo))                    
            if 'F' in data_load['sexoDeHijoPatrocinado1'] and 'F' in data_load['sexoDeHijoPatrocinado2'] and 'F' in data_load['sexoDeHijoPatrocinado3'] and 'F' in data_load['sexoDeHijoPatrocinado4']:
                parrafo6.add_run(" sus hijas ")
            else:
                parrafo6.add_run(" sus hijos ")            
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado1'])
            parrafo6.add_run(", ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado2'])
            parrafo6.add_run(", ")                        
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado3'])
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado4'])
            parrafo6.add_run(".")

        elif data_load['hijosPatrocinado']==5:
            parrafo6.add_run(str(numeroParrafo))                    
            if 'F' in data_load['sexoDeHijoPatrocinado1'] and 'F' in data_load['sexoDeHijoPatrocinado2'] and 'F' in data_load['sexoDeHijoPatrocinado3'] and 'F' in data_load['sexoDeHijoPatrocinado4'] and 'F' in data_load['sexoDeHijoPatrocinado5']:
                parrafo6.add_run(" sus hijas ")
            else:
                parrafo6.add_run(" sus hijos ")            
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado1'])
            parrafo6.add_run(", ")            
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado2'])
            parrafo6.add_run(", ")            
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado3'])
            parrafo6.add_run(", ")            
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado4'])     
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado5'])
            parrafo6.add_run(", mientras que por parte de ")        
            parrafo6.add_run(data_load['honorificoContraparte'])
            parrafo6.add_run(" ")
            parrafo6.add_run(data_load['nombreContraparte'])
            parrafo6.add_run(" ")
            parrafo6.add_run(data_load['apellidoPaternoContraparte'])
            parrafo6.add_run(" ")
            parrafo6.add_run(data_load['apellidoMaternoContraparte'])         
        if data_load['hijosContraparte']==0:
                parrafo6.add_run("no han nacido hijos.")
        elif data_load['hijosContraparte']==1:
            if 'F' in data_load['sexoDeHijoContraparte1']:
                parrafo6.add_run(" ha nacido su hija ")
            elif 'M' in data_load['sexoDeHijoContraparte1']:
                parrafo6.add_run(" ha nacido su hijo ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte1'])
            parrafo6.add_run(".")
        elif data_load['hijosContraparte']==2:
            parrafo6.add_run(str(numeroParrafo))                    
            if 'F' in data_load['sexoDeHijoContraparte1'] and 'F' in data_load['sexoDeHijoContraparte2']:
                parrafo6.add_run(" han nacido sus hijas ")
            else:
                parrafo6.add_run(" han nacido sus hijos ")            
            parrafo6.add_run(data_load['nombreDeHijoContraparte1'])
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte2'])
            parrafo6.add_run(".")

        elif data_load['hijosContraparte']==3:
            parrafo6.add_run(str(numeroParrafo))                    
            if 'F' in data_load['sexoDeHijoContraparte1'] and 'F' in data_load['sexoDeHijoContraparte2'] and 'F' in data_load['sexoDeHijoContraparte3']:
                parrafo6.add_run(" han nacido sus hijas ")
            else:
                parrafo6.add_run(" han nacido sus hijos ")            
            parrafo6.add_run(data_load['nombreDeHijoContraparte1'])
            parrafo6.add_run(", ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte2'])            
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte3'])
            parrafo6.add_run(".")

        elif data_load['hijosContraparte']==4:
            parrafo6.add_run(str(numeroParrafo))                    
            if 'F' in data_load['sexoDeHijoContraparte1'] and 'F' in data_load['sexoDeHijoContraparte2'] and 'F' in data_load['sexoDeHijoContraparte3'] and 'F' in data_load['sexoDeHijoContraparte4']:
                parrafo6.add_run(" han nacido sus hijas ")
            else:
                parrafo6.add_run(" han nacido sus hijos ")            
            parrafo6.add_run(data_load['nombreDeHijoContraparte1'])
            parrafo6.add_run(", ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte2'])
            parrafo6.add_run(", ")                        
            parrafo6.add_run(data_load['nombreDeHijoContraparte3'])
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte4'])
            parrafo6.add_run(".")

        elif data_load['hijosContraparte']==5:
            parrafo6.add_run(str(numeroParrafo))                    
            if 'F' in data_load['sexoDeHijoContraparte1'] and 'F' in data_load['sexoDeHijoContraparte2'] and 'F' in data_load['sexoDeHijoContraparte3'] and 'F' in data_load['sexoDeHijoContraparte4'] and 'F' in data_load['sexoDeHijoContraparte5']:
                parrafo6.add_run(" han nacido sus hijas ")
            else:
                parrafo6.add_run(" han nacido sus hijos ") 
            parrafo6.add_run(data_load['nombreDeHijoContraparte1'])
            parrafo6.add_run(", ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte2'])
            parrafo6.add_run(", ")                        
            parrafo6.add_run(data_load['nombreDeHijoContraparte3'])
            parrafo6.add_run(", ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte4'])
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte5'])            
            parrafo6.add_run(".")            

    
        formatoParrafo6=parrafo6.paragraph_format
        formatoParrafo6.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY        
    elif data_load['nuevaRelacionPatrocinado']=='Si':
        numeroParrafo+=1
        parrafo6=document.add_paragraph()
        parrafo6.add_run(str(numeroParrafo))
        parrafo6.add_run(".- Con posterioridad a nuestra separación ")
        parrafo6.add_run(data_load['nombrePatrocinado'])
        parrafo6.add_run(" ")
        parrafo6.add_run(data_load['apellidoPaternoPatrocinado'])
        parrafo6.add_run(" ")
        parrafo6.add_run(data_load['apellidoMaternoPatrocinado'])
        parrafo6.add_run(" ha iniciado una nueva relación que ha perdurado hasta el día de hoy, la cual se extiende por más de ")
        if int(data_load['añosRelacionPatrocinado'])==1:
            parrafo6.add_run('un año')
        elif int(data_load['añosRelacionPatrocinado'])>1:
            parrafo6.add_run(data_load['añosRelacionPatrocinado'])
            parrafo6.add_run(" años")
        if data_load['hijosPatrocinado']==0:
            parrafo6.add_run('.')
        elif data_load['hijosPatrocinado']==1:
            if 'F' in data_load['sexoDeHijoPatrocinado1']:
                parrafo6.add_run(",producto de la cual nació su hija ")
            elif 'M' in data_load['sexoDeHijoPatrocinado1']:
                parrafo6.add_run(", producto de la cual nació su hijo ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado1'])
            parrafo6.add_run(".")

        elif data_load['hijosPatrocinado']==2:
            if 'F' in data_load['sexoDeHijoPatrocinado1'] and 'F' in data_load['sexoDeHijoPatrocinado2']:
                parrafo6.add_run(", producto de la cual nacieron sus hijas ")
            else:
                parrafo6.add_run(", producto de la cual nacieron sus hijos ")   
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado1'])
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado2'])
            parrafo6.add_run(".")

        elif data_load['hijosPatrocinado']==3:
            if 'F' in data_load['sexoDeHijoPatrocinado1'] and 'F' in data_load['sexoDeHijoPatrocinado2'] and 'F' in data_load['sexoDeHijoPatrocinado3']:
                parrafo6.add_run(", producto de la cual nacieron sus hijas ")
            else:
                parrafo6.add_run(", producto de la cual nacieron sus hijos ") 
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado1'])
            parrafo6.add_run(", ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado2'])            
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado3'])
            parrafo6.add_run(".")

        elif data_load['hijosPatrocinado']==4:
            parrafo6.add_run(str(numeroParrafo))                    
            if 'F' in data_load['sexoDeHijoPatrocinado1'] and 'F' in data_load['sexoDeHijoPatrocinado2'] and 'F' in data_load['sexoDeHijoPatrocinado3'] and 'F' in data_load['sexoDeHijoPatrocinado4']:
                parrafo6.add_run(", producto de la cual nacieron sus hijas ")
            else:
                parrafo6.add_run(", producto de la cual nacieron sus hijos ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado1'])
            parrafo6.add_run(", ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado2'])
            parrafo6.add_run(", ")                        
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado3'])
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado4'])
            parrafo6.add_run(".")

        elif data_load['hijosPatrocinado']==5:                 
            if 'F' in data_load['sexoDeHijoPatrocinado1'] and 'F' in data_load['sexoDeHijoPatrocinado2'] and 'F' in data_load['sexoDeHijoPatrocinado3'] and 'F' in data_load['sexoDeHijoPatrocinado4'] and 'F' in data_load['sexoDeHijoPatrocinado5']:
                parrafo6.add_run(", producto de la cual nacieron sus hijas ")
            else:
                parrafo6.add_run(", producto de la cual nacieron sus hijos ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado1'])
            parrafo6.add_run(", ")            
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado2'])
            parrafo6.add_run(", ")            
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado3'])
            parrafo6.add_run(", ")            
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado4'])     
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoPatrocinado5'])
            parrafo6.add_run(".")            

             
        formatoParrafo6=parrafo6.paragraph_format
        formatoParrafo6.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY        
    elif data_load['nuevaRelacionContraparte']=='Si':
        numeroParrafo+=1
        parrafo6=document.add_paragraph()
        parrafo6.add_run(str(numeroParrafo))
        parrafo6.add_run(".- Con posterioridad a nuestra separación ")
        parrafo6.add_run(data_load['honorificoContraparte'])
        parrafo6.add_run(" ")                    
        parrafo6.add_run(data_load['nombreContraparte'])
        parrafo6.add_run(" ")
        parrafo6.add_run(data_load['apellidoPaternoContraparte'])
        parrafo6.add_run(" ")
        parrafo6.add_run(data_load['apellidoMaternoContraparte']) 
        parrafo6.add_run(" ha iniciado una nueva relación que ha perdurado hasta el día de hoy, la cual se extiende por más de ")
        parrafo6.add_run(data_load['añosRelacionContraparte'])
        parrafo6.add_run(" años")
        if data_load['hijosContraparte']==0:
            parrafo6.add_run('.')
            
        elif data_load['hijosContraparte']==1:
            if 'F' in data_load['sexoDeHijoContraparte1']:
                parrafo6.add_run(",producto de la cual nació su hija ")
            elif 'M' in data_load['sexoDeHijoContraparte1']:
                parrafo6.add_run(", producto de la cual nació su hijo ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte1'])
            parrafo6.add_run(".")

        elif data_load['hijosContraparte']==2:
            if 'F' in data_load['sexoDeHijoContraparte1'] and 'F' in data_load['sexoDeHijoContraparte2']:
                parrafo6.add_run(", producto de la cual nacieron sus hijas ")
            else:
                parrafo6.add_run(", producto de la cual nacieron sus hijos ") 
            parrafo6.add_run(data_load['nombreDeHijoContraparte1'])
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte2'])
            parrafo6.add_run(".")

        elif data_load['hijosContraparte']==3:                 
            if 'F' in data_load['sexoDeHijoContraparte1'] and 'F' in data_load['sexoDeHijoContraparte2'] and 'F' in data_load['sexoDeHijoContraparte3']:
                parrafo6.add_run(", producto de la cual nacieron sus hijas ")
            else:
                parrafo6.add_run(", producto de la cual nacieron sus hijos ")            
            parrafo6.add_run(data_load['nombreDeHijoContraparte1'])
            parrafo6.add_run(", ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte2'])            
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte3'])
            parrafo6.add_run(".")

        elif data_load['hijosContraparte']==4:                    
            if 'F' in data_load['sexoDeHijoContraparte1'] and 'F' in data_load['sexoDeHijoContraparte2'] and 'F' in data_load['sexoDeHijoContraparte3'] and 'F' in data_load['sexoDeHijoContraparte4']:
                parrafo6.add_run(", producto de la cual nacieron sus hijas ")
            else:
                parrafo6.add_run(", producto de la cual nacieron sus hijos ") 
            parrafo6.add_run(data_load['nombreDeHijoContraparte1'])
            parrafo6.add_run(", ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte2'])
            parrafo6.add_run(", ")                        
            parrafo6.add_run(data_load['nombreDeHijoContraparte3'])
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte4'])
            parrafo6.add_run(".")

        elif data_load['hijosContraparte']==5:                   
            if 'F' in data_load['sexoDeHijoContraparte1'] and 'F' in data_load['sexoDeHijoContraparte2'] and 'F' in data_load['sexoDeHijoContraparte3'] and 'F' in data_load['sexoDeHijoContraparte4'] and 'F' in data_load['sexoDeHijoContraparte5']:
                parrafo6.add_run(", producto de la cual nacieron sus hijas ")
            else:
                parrafo6.add_run(", producto de la cual nacieron sus hijos ") 
            parrafo6.add_run(data_load['nombreDeHijoContraparte1'])
            parrafo6.add_run(", ")            
            parrafo6.add_run(data_load['nombreDeHijoContraparte2'])
            parrafo6.add_run(", ")            
            parrafo6.add_run(data_load['nombreDeHijoContraparte3'])
            parrafo6.add_run(", ")            
            parrafo6.add_run(data_load['nombreDeHijoContraparte4'])     
            parrafo6.add_run(" y ")
            parrafo6.add_run(data_load['nombreDeHijoContraparte5'])
            parrafo6.add_run(".")   
        formatoParrafo6=parrafo6.paragraph_format
        formatoParrafo6.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    numeroParrafo+=1
    parrafo7=document.add_paragraph()
    parrafo7.add_run(str(numeroParrafo))
    parrafo7.add_run(".- De esta forma S.S., se cumple así el requisito que nuestra legislación establece para poder solicitar el divorcio de manera unilateral por cese de convivencia. Por tal razón, mi solicitud tiene como objeto la disolución de nuestro matrimonio, y de este modo generar una coherencia entre la realidad fáctica y la realidad jurídica de nuestra relación matrimonial.")

    formatoParrafo7=parrafo7.paragraph_format
    formatoParrafo7.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    numeroParrafo+=1
    parrafo8=document.add_paragraph()
    parrafo8.add_run(str(numeroParrafo))
    if data_load['pideCompensacionEconomica']=='Si':
        parrafo8.add_run(".- De la misma forma, hago presente a S.S. que, habiendo sido oportunamente informad")
        parrafo8.add_run(data_load['generoPatrocinado'])
        parrafo8.add_run(" de la institución de la Compensación Económica, he decidido solicitar se me otorgue, en atención a que cumplo con los requisitos exigidos por el legislador para ello, como se acreditara oportunamente.")


    elif data_load['pideCompensacionEconomica']=='No':
        parrafo8.add_run(".- De la misma forma, hago presente a S.S., que, habiendo sido informad")
        parrafo8.add_run(data_load['generoPatrocinado'])
        parrafo8.add_run(" acerca de la institucion de la Compensacion Economica, declaro renunciar expresamente a ella.")

    formatoParrafo8=parrafo8.paragraph_format
    formatoParrafo8.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_paragraph()
    derecho="EL DERECHO:"
    parrafo9=document.add_paragraph()
    parrafo9.add_run(derecho).bold=True
    document.add_paragraph()


    numeroParrafo=0

    numeroParrafo+=1

    parrafo10=document.add_paragraph()
    parrafo10.add_run(str(numeroParrafo))
    parrafo10.add_run(".- La Ley 19.947 sobre Matrimonio Civil dispone en su artículo 42 número 4 que; “el matrimonio termina: Por sentencia firme de divorcio”, cuestión que solicito por este acto.")

    formatoParrafo10=parrafo10.paragraph_format
    formatoParrafo10.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


    numeroParrafo+=1
    parrafo10=document.add_paragraph()
    parrafo10.add_run(str(numeroParrafo))
    parrafo10.add_run('.- Conforme a lo establecido en el artículo 55 inciso 3 de la Ley Nº 19.947,  el cual dispone que “habrá lugar también al divorcio cuando se verifique un cese efectivo de la convivencia conyugal durante el transcurso de, a lo menos, tres años”. Solicito a S.S. el término del vínculo legal de nuestro matrimonio por cumplirse el cese efectivo de la convivencia con mi cónyuge, al reunirse los requisitos establecidos en la ley para tales efectos. El caso de autos cumple con todos los requisitos de forma y fondo exigidos, por cuanto, tal y como se acreditará en la etapa procesal correspondiente, nuestra vida en común cesó hace más de tres años.')

    formatoParrafo10=parrafo10.paragraph_format
    formatoParrafo10.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


#QUE ESTE NUMERAL VARIE SEGUN SI HAY O NO HIJOS MENORES DE EDAD
    numeroParrafo+=1
    parrafo10=document.add_paragraph()
    parrafo10.add_run(str(numeroParrafo))
    parrafo10.add_run(".-El artículo 61 de la Ley de Matrimonio Civil establece la institución de la Compensación Económica entre los cónyuges en casos en que el matrimonio termine ya sea por nulidad o por divorcio como es en mi caso. Dicha disposición legal señala: “Si, como consecuencia de haberse dedicado al cuidado de los hijos o a las labores propias del hogar común, uno de los cónyuges no pudo desarrollar una actividad remunerada o lucrativa durante el matrimonio, o lo hizo en menor medida de lo que podía y quería, tendrá derecho a que, cuando se produzca el divorcio o se declare la nulidad del matrimonio, se le compense el menoscabo económico sufrido por esta causa.”.")
    if data_load['pideCompensacionEconomica']=="Si":
        parrafo10.add_run(" Dichos presupuestos se cumplen en el presente caso, por cuanto .")
        parrafo10.add_run(" Esta compensacion se considerara saldada con .")
    elif data_load['pideCompensacionEconomica']=="No":
        parrafo10.add_run("En mi caso y en razón del artículo 12 de nuestro Código Civil he decidido renunciar a la facultad de exigir compensación económica en razón de que dicho precepto legal establece la facultad de renunciar a aquellos derechos que sólo miran el interés individual y su renuncia no se encuentra prohibida por nuestra legislación ni perjudica derechos de terceros, disposición absolutamente aplicable mi caso.") 

    formatoParrafo10=parrafo10.paragraph_format
    formatoParrafo10.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    formatoParrafo10=parrafo10.paragraph_format
    formatoParrafo10.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    if int(data_load['añoDeMatrimonio'])>2004:
        numeroParrafo+=1
        parrafo10=document.add_paragraph()
        parrafo10.add_run(str(numeroParrafo))
        parrafo10.add_run(".- Por tratarse de un matrimonio celebrado con posterioridad al 17 de noviembre del año 2004, resultan aplicables los artículos 22 y 25 de la ley de matrimonio civil, en los cuales se acotan las formas válidas de acreditar el cese de convivencia.")
    elif int(data_load['añoDeMatrimonio'])==2004:
        if mesMatrimonio=="Diciembre":
            numeroParrafo+=1
            parrafo10=document.add_paragraph()
            parrafo10.add_run(str(numeroParrafo))
            parrafo10.add_run(".- Por tratarse de un matrimonio celebrado con posterioridad al 17 de noviembre del año 2004, resultan aplicables los artículos 22 y 25 de la ley de matrimonio civil, en los cuales se acotan las formas válidas de acreditar el cese de convivencia.")
        elif mesMatrimonio=="Noviembre":
            if int(diaMatrimonio)>17:
                numeroParrafo+=1
                parrafo10=document.add_paragraph()
                parrafo10.add_run(str(numeroParrafo))
                parrafo10.add_run(".- Por tratarse de un matrimonio celebrado con posterioridad al 17 de noviembre del año 2004, resultan aplicables los artículos 22 y 25 de la ley de matrimonio civil, en los cuales se acotan las formas válidas de acreditar el cese de convivencia.")

    formatoParrafo10=parrafo10.paragraph_format
    formatoParrafo10.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    
    parrafo12=document.add_paragraph()
    parrafo12.add_run("POR TANTO,").bold=True
    parrafo12.add_run(" en virtud de los hechos y el derecho expuestos y  de  los artículos 42 Nº 4, 55 inciso 3 de la Ley 19.947 sobre Matrimonio Civil, los artículos 55 y siguientes de la Ley 19.968 sobre Tribunales de Familia y demás normas pertinentes:")
    
    parrafo12=document.add_paragraph()
    parrafo12.add_run("SOLICITO A S.S.:").bold=True
    parrafo12.add_run(" se sirva tener por interpuesta demanda de divorcio unilateral en contra de ")
    parrafo12.add_run(data_load['honorificoContraparte'])
    parrafo12.add_run(" ")    
    parrafo12.add_run(data_load['nombreContraparte'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoPaternoContraparte'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoMaternoContraparte'])
    parrafo12.add_run(", ya debidamente individualizad")
    parrafo12.add_run(data_load['generoContraparte'])                      
    parrafo12.add_run(", acogerla a tramitación en todas sus partes y, en definitiva, decretar el término de nuestro vínculo matrimonial y ordenar las correspondientes subinscripciones")
    if data_load['pideCompensacionEconomica']=="Si":
        parrafo12.add_run(", asi como el pago de la compensacion economica ya expuesta")
    parrafo12.add_run(".")
                          
    #DOCUMENTOS
    parrafo12=document.add_paragraph()
    parrafo12.add_run("PRIMER OTROSÍ:").bold=True
    parrafo12.add_run(" Que, por este acto, acompaño los siguientes documentos:")
    document.add_paragraph("")

    parrafo12=document.add_paragraph()
    parrafo12.add_run("    - Certificado de matrimonio entre ")
    parrafo12.add_run(data_load['honorificoPatrocinado'])
    parrafo12.add_run(" ")                    
    parrafo12.add_run(data_load['nombrePatrocinado'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoPaternoPatrocinado'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoMaternoPatrocinado'])
    parrafo12.add_run(" y ")                    
    parrafo12.add_run(data_load['honorificoContraparte'])
    parrafo12.add_run(" ")                    
    parrafo12.add_run(data_load['nombreContraparte'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoPaternoContraparte'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoMaternoContraparte'])
    parrafo12.add_run(".")                    
                
    if 'a' in data_load['nombreDeHijoEntreParte1'] or 'e' in data_load['nombreDeHijoEntreParte1'] or 'i' in data_load['nombreDeHijoEntreParte1'] or 'o' in data_load['nombreDeHijoEntreParte1'] or 'u' in data_load['nombreDeHijoEntreParte1']:
        parrafo12=document.add_paragraph()                    
        parrafo12.add_run("    - Certificado de nacimiento de ")
        parrafo12.add_run(data_load['nombreDeHijoEntreParte1'])
        parrafo12.add_run(".")                            
    if 'a' in data_load['nombreDeHijoEntreParte2'] or 'e' in data_load['nombreDeHijoEntreParte2'] or 'i' in data_load['nombreDeHijoEntreParte2'] or 'o' in data_load['nombreDeHijoEntreParte2'] or 'u' in data_load['nombreDeHijoEntreParte2']:
        parrafo12=document.add_paragraph()                                        
        parrafo12.add_run("    - Certificado de nacimiento de ")
        parrafo12.add_run(data_load['nombreDeHijoEntreParte2'])
        parrafo12.add_run(".")                       
    if 'a' in data_load['nombreDeHijoEntreParte3'] or 'e' in data_load['nombreDeHijoEntreParte3'] or 'i' in data_load['nombreDeHijoEntreParte3'] or 'o' in data_load['nombreDeHijoEntreParte3'] or 'u' in data_load['nombreDeHijoEntreParte3']:
        parrafo12=document.add_paragraph()                                        
        parrafo12.add_run("    - Certificado de nacimiento de ")
        parrafo12.add_run(data_load['nombreDeHijoEntreParte3'])
        parrafo12.add_run(".")                     
    if 'a' in data_load['nombreDeHijoEntreParte4'] or 'e' in data_load['nombreDeHijoEntreParte4'] or 'i' in data_load['nombreDeHijoEntreParte4'] or 'o' in data_load['nombreDeHijoEntreParte4'] or 'u' in data_load['nombreDeHijoEntreParte4']:
        parrafo12=document.add_paragraph()                                        
        parrafo12.add_run("    - Certificado de nacimiento de ")
        parrafo12.add_run(data_load['nombreDeHijoEntreParte4'])
        parrafo12.add_run(".")                     
    if 'a' in data_load['nombreDeHijoEntreParte5'] or 'e' in data_load['nombreDeHijoEntreParte5'] or 'i' in data_load['nombreDeHijoEntreParte5'] or 'o' in data_load['nombreDeHijoEntreParte5'] or 'u' in data_load['nombreDeHijoEntreParte5']:
        parrafo12=document.add_paragraph()                                        
        parrafo12.add_run("    - Certificado de nacimiento de ")
        parrafo12.add_run(data_load['nombreDeHijoEntreParte5'])
        parrafo12.add_run(".")                     

    parrafo12=document.add_paragraph()
    parrafo12.add_run("    - Certificado de residencia de ")
    parrafo12.add_run(data_load['honorificoPatrocinado'])
    parrafo12.add_run(" ")                    
    parrafo12.add_run(data_load['nombrePatrocinado'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoPaternoPatrocinado'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoMaternoPatrocinado'])
    parrafo12.add_run(".")
                     
                          
    parrafo12=document.add_paragraph()
    parrafo12.add_run("POR TANTO,").bold=True
    parrafo12=document.add_paragraph()
    parrafo12.add_run("SOLICITO A S.S.: ").bold=True
    parrafo12.add_run("se sirva tenerlos por acompañados.")

    parrafo12=document.add_paragraph()
    parrafo12.add_run("SEGUNDO OTROSÍ:").bold=True
    parrafo12.add_run(" Ruego a U.S., autorizarme para litigar en forma electrónica y ser notificad")
    parrafo12.add_run(data_load['generoPatrocinado'])
    parrafo12.add_run(" de las resoluciones que se dicten en la presente causa a través del correo electrónico ")
    parrafo12.add_run(data_load['emailAbogadoPatrocinado'])
    parrafo12.add_run(".")
    
    parrafo12=document.add_paragraph()
    parrafo12.add_run("TERCER OTROSÍ:").bold=True
    parrafo12.add_run(" A U.S., pido  tener presente que ")
    parrafo12.add_run(" designo abogad")
    parrafo12.add_run(data_load['generoAbogadoPatrocinado'])
    parrafo12.add_run(" patrocinante y confiero poder a ")
    parrafo12.add_run(data_load['honorificoAbogadoPatrocinado'])
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['nombreAbogadoPatrocinado'].upper()).bold=True
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoPaternoAbogadoPatrocinado'].upper()).bold=True
    parrafo12.add_run(" ")
    parrafo12.add_run(data_load['apellidoMaternoAbogadoPatrocinado'].upper()).bold=True
    parrafo12.add_run(", cédula de identidad Nº ")
    parrafo12.add_run(data_load['rutAbogadoPatrocinado'])
    parrafo12.add_run(", abogad")
    parrafo12.add_run(data_load['generoAbogadoPatrocinado'])
    parrafo12.add_run(" habilitad")
    parrafo12.add_run(data_load['generoAbogadoPatrocinado'])
    parrafo12.add_run(" para el ejercicio de la profesion, con domicilio en ")
    parrafo12.add_run(data_load['domicilioAbogadoPatrocinado'])
    parrafo12.add_run(" número ")
    parrafo12.add_run(data_load['numeroDomicilioAbogadoPatrocinado']) 
    parrafo12.add_run(" comuna de ")
    parrafo12.add_run(data_load['comunaAbogadoPatrocinado'])                    
    if data_load['facultadesAbogadoPatrocinado'] == 'Sí':
        parrafo12.add_run(' para que ')
        parrafo12.add_run(data_load['conectorPatrocinado'])
        parrafo12.add_run(' represente con las facultades de ambos incisos del artículo 7° del Código de Procedimiento Civil, en especial, las de percibir, avenir y transigir,')    
    parrafo12.add_run(" que firma en señal de aceptación")
    if 'a' in data_load['nombreApoderadoPatrocinado'] or 'e' in data_load['nombreApoderadoPatrocinado'] or 'i' in data_load['nombreApoderadoPatrocinado'] or 'o' in data_load['nombreApoderadoPatrocinado'] or 'u' in data_load['nombreApoderadoPatrocinado']:
        parrafo12.add_run("; y también ")
        parrafo12.add_run(conector2ApoderadoPatrocinado)
        parrafo12.add_run(" ")
        parrafo12.add_run(profesionApoderadoPatrocinado)
        parrafo12.add_run(" ")
        parrafo12.add_run(data_load['nombreApoderadoPatrocinado'].upper())
        parrafo12.add_run(" ")
        parrafo12.add_run(data_load['apellidoPaternoApoderadoPatrocinado'].upper())
        parrafo12.add_run(" ")
        parrafo12.add_run(data_load['apellidoMaternoApoderadoPatrocinado'].upper())                                   
        parrafo12.add_run(", cédula de identidad Nº ")
        parrafo12.add_run(data_load['rutApoderadoPatrocinado'])
        parrafo12.add_run(", del mismo domicilio y forma de notificación")
        if data_load['facultadesApoderadoPatrocinado'] == 'Sí':
            parrafo12.add_run(', para que ')
            parrafo12.add_run(data_load['conectorPatrocinado'])
            parrafo12.add_run(' represente con las facultades de ambos incisos del artículo 7° del Código de Procedimiento Civil, en especial, las de percibir, avenir y transigir,')
        else:
            parrafo12.add_run('.')        
    else:
        parrafo12.add_run(".")    


    document.add_paragraph("")

    parrafo12=document.add_paragraph()
    parrafo12.add_run("POR TANTO,").bold=True

    parrafo12=document.add_paragraph()
    parrafo12.add_run("SOLICITO A S.S.: ").bold=True
    parrafo12.add_run("Se sirva tener presente patrocinio y poder conferido.")




    formatoParrafo12=parrafo10.paragraph_format
    formatoParrafo12.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save('DivUniCese'+data_load['apellidoPaternoPatrocinado']+data_load['apellidoPaternoContraparte']+'.docx')

    print("I am a zombie, ZOMBIE!!")
