from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Length
from docx.text.run import Font
from docx.table import Table
from tkinter import *
from tkinter.filedialog import *
from imr import ingresoMinimo
import json


def acuerdoCompletoSimple():
    with open(filedialog.askopenfilename()) as fp:
            data_load=json.load(fp)
    
    #DOCUMENTO PROPIAMENTE TAL

    document=Document("documentoBase.docx")
    Font.name = 'Century Gothic'
    Font.size = Pt(16)

    document.add_paragraph('')
    document.add_paragraph('')
    encabezado1=document.add_paragraph()
    encabezado1.add_run('ACUERDO COMPLETO Y SUFICIENTE SIMPLE').bold=True


    
    encabezado2=document.add_paragraph()
    encabezado2.add_run("S.J. DE FAMILIA DE SANTIAGO (").bold=True
    encabezado2.add_run(data_load['juzgado']).bold=True
    encabezado2.add_run(")").bold=True
    formatoEncabezado2=encabezado2.paragraph_format
    formatoEncabezado2.alignment=WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")


    parrafo1=document.add_paragraph()

    parrafo1.add_run("    ")
    parrafo1.add_run(data_load['nombrePatrocinado'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoPatrocinado'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoPatrocinado'].upper()).bold=True
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['estadoCivilPatrocinado'])
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['nacionalidadPatrocinado'])
    parrafo1.add_run(", cédula de identidad Nº ")
    parrafo1.add_run(data_load['rutPatrocinado'])
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['profesionPatrocinado'])
    parrafo1.add_run(", domiciliad")
    parrafo1.add_run(data_load['generoPatrocinado'])
    parrafo1.add_run(" en ")
    parrafo1.add_run(data_load['domicilioPatrocinado'])
    parrafo1.add_run(" número ")
    parrafo1.add_run(data_load['numeroDomicilioPatrocinado'])    
    parrafo1.add_run(", comuna de ")
    parrafo1.add_run(data_load['comunaPatrocinado'])
    parrafo1.add_run("; y ")

    parrafo1.add_run(data_load['nombreContraparte'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoPaternoContraparte'].upper()).bold=True
    parrafo1.add_run(" ")
    parrafo1.add_run(data_load['apellidoMaternoContraparte'].upper()).bold=True
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['estadoCivilContraparte'])
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['nacionalidadContraparte'])
    parrafo1.add_run(", cédula de identidad Nº ")
    parrafo1.add_run(data_load['rutContraparte'])
    parrafo1.add_run(", ")
    parrafo1.add_run(data_load['profesionContraparte'])
    parrafo1.add_run(", domiciliad")
    parrafo1.add_run(data_load['generoContraparte'])
    parrafo1.add_run(" en ")
    parrafo1.add_run(data_load['domicilioContraparte'])
    parrafo1.add_run(" número ")
    parrafo1.add_run(data_load['numeroDomicilioContraparte'])     
    parrafo1.add_run(", comuna de ")
    parrafo1.add_run(data_load['comunaContraparte'])
    parrafo1.add_run("; a U.S., respetuosamente decimos:")

    formatoParrafo1=parrafo1.paragraph_format
    formatoParrafo1.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    formatoParrafo1.line_spacing=1.5

    parrafo3=document.add_paragraph()
    parrafo3.add_run(" Que el dia ")
    parrafo3.add_run(data_load['diaDeMatrimonio'])
    parrafo3.add_run(" de ")
    parrafo3.add_run(data_load['mesDeMatrimonio'])
    parrafo3.add_run(" de ")
    parrafo3.add_run(data_load['añoDeMatrimonio'])
    parrafo3.add_run(" contrajimos matrimonio ante Oficial de Registro Civil de la circunscripción de ")
    parrafo3.add_run(data_load['comunaDeMatrimonio'])
    parrafo3.add_run(", el que fue anotado bajo la inscripción Nº ")
    parrafo3.add_run(data_load['numeroInscripcionMatrimonio'])
    parrafo3.add_run(" de dicho año pactando el régimen matrimonial de ")
    parrafo3.add_run(data_load['regimenMatrimonial'].lower())
    parrafo3.add_run(".")

    formatoParrafo3=parrafo3.paragraph_format
    formatoParrafo3.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


    if data_load['hijosEntrePartes']==0:
        parrafo3.add_run(" Producto de nuestra unión no nacieron hijos.")
    elif data_load['hijosEntrePartes']==1:
        if 'F' in data_load['sexoDeHijoEntreParte1']:
            parrafo3.add_run(" Producto de nuestra union nació nuestra hija ")
        elif 'M' in data_load['sexoDeHijoEntreParte1']:
            parrafo3.add_run(" Producto de nuestra unión nació nuestro hijo ")
        parrafo3.add_run(data_load['nombreDeHijoEntreParte1'])
        parrafo3.add_run(", RUT ")
        parrafo3.add_run(data_load['rutDeHijoEntreParte1'])
        parrafo3.add_run(", de actuales ")
        parrafo3.add_run(data_load['edadDeHijoEntreParte1'])
        parrafo3.add_run(" años de edad.")
    elif data_load['hijosEntrePartes']>1:                  
        parrafo3.add_run(" Producto de nuestra unión nacieron ")
        if data_load['hijosEntrePartes'] == 2:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2']:
                parrafo3.add_run("nuestras hijas:")
            else:
                parrafo3.add_run("nuestros hijos:")
        elif data_load['hijosEntrePartes'] == 3:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3']:
                parrafo3.add_run("nuestras hijas:")
            else:
                parrafo3.add_run("nuestros hijos:")
        elif data_load['hijosEntrePartes'] == 4:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3'] and 'F' in data_load['sexoDeHijoEntreParte4']:
                parrafo3.add_run("nuestras hijas:")
            else:
                parrafo3.add_run("nuestros hijos:")            
        elif data_load['hijosEntrePartes'] == 5:
            if 'F' in data_load['sexoDeHijoEntreParte1'] and 'F' in data_load['sexoDeHijoEntreParte2'] and 'F' in data_load['sexoDeHijoEntreParte3'] and 'F' in data_load['sexoDeHijoEntreParte4'] and 'F' in data_load['sexoDeHijoEntreParte5']:
                parrafo3.add_run("nuestras hijas:")
            else:
                parrafo3.add_run("nuestros hijos:")            
        parrafo3=document.add_paragraph()
        parrafo3.add_run("- ")
        parrafo3.add_run(data_load['nombreDeHijoEntreParte1'])
        parrafo3.add_run(", RUT ")
        parrafo3.add_run(data_load['rutDeHijoEntreParte1'])
        parrafo3.add_run(", de actuales ")
        parrafo3.add_run(data_load['edadDeHijoEntreParte1'])
        parrafo3.add_run(" años de edad.")
        parrafo3=document.add_paragraph()
        parrafo3.add_run("- ")
        parrafo3.add_run(data_load['nombreDeHijoEntreParte2'])
        parrafo3.add_run(", RUT ")
        parrafo3.add_run(data_load['rutDeHijoEntreParte2'])
        parrafo3.add_run(", de actuales ")
        parrafo3.add_run(data_load['edadDeHijoEntreParte2'])
        parrafo3.add_run(" años de edad.")
        if 'a' in data_load['nombreDeHijoEntreParte3'] or 'e' in data_load['nombreDeHijoEntreParte3'] or 'i' in data_load['nombreDeHijoEntreParte3'] or 'o' in data_load['nombreDeHijoEntreParte3'] or 'u' in data_load['nombreDeHijoEntreParte3']:
            parrafo3=document.add_paragraph()
            parrafo3.add_run("- ")
            parrafo3.add_run(data_load['nombreDeHijoEntreParte3'])
            parrafo3.add_run(", RUT ")
            parrafo3.add_run(data_load['rutDeHijoEntreParte3'])
            parrafo3.add_run(", de actuales ")
            parrafo3.add_run(data_load['edadDeHijoEntreParte3'])
            parrafo3.add_run(" años de edad.")
        if 'a' in data_load['nombreDeHijoEntreParte4'] or 'e' in data_load['nombreDeHijoEntreParte4'] or 'i' in data_load['nombreDeHijoEntreParte4'] or 'o' in data_load['nombreDeHijoEntreParte4'] or 'u' in data_load['nombreDeHijoEntreParte4']:
            parrafo3=document.add_paragraph()
            parrafo3.add_run("- ")
            parrafo3.add_run(data_load['nombreDeHijoEntreParte4'])
            parrafo3.add_run(", RUT ")
            parrafo3.add_run(data_load['rutDeHijoEntreParte4'])
            parrafo3.add_run(", de actuales ")
            parrafo3.add_run(data_load['edadDeHijoEntreParte4'])
            parrafo3.add_run(" años de edad.")
        if 'a' in data_load['nombreDeHijoEntreParte5'] or 'e' in data_load['nombreDeHijoEntreParte5'] or 'i' in data_load['nombreDeHijoEntreParte5'] or 'o' in data_load['nombreDeHijoEntreParte5'] or 'u' in data_load['nombreDeHijoEntreParte5']:
            parrafo3=document.add_paragraph()
            parrafo3.add_run("- ")
            parrafo3.add_run(data_load['nombreDeHijoEntreParte5'])
            parrafo3.add_run(", RUT ")
            parrafo3.add_run(data_load['rutDeHijoEntreParte5'])
            parrafo3.add_run(", de actuales ")
            parrafo3.add_run(data_load['edadDeHijoEntreParte5'])
            parrafo3.add_run(" años de edad.")                    
    formatoparrafo3=parrafo3.paragraph_format
    formatoparrafo3.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY



    parrafo5=document.add_paragraph()
    parrafo5.add_run(" Que acompañamos el siguiente acuerdo para que S.S. lo califique de completo y suficiente, autorice y apruebe ")
    parrafo5.add_run("en la respectiva audiencia, en esta causa de divorcio por cese de la convivencia de común acuerdo.")

    parrafo6=document.add_paragraph()
    parrafo6.add_run(' En efecto, ambas partes hemos acordado regular todas las materias relevantes y ')
    parrafo6.add_run('que la Ley Nº 19.947 en su artículo 55 en relación con el artículo 21 del mismo cuerpo normativo exige, para ')
    parrafo6.add_run('que S.S. acoja la presente demanda de divorcio por cese de la convivencia de común acuerdo en la ')
    parrafo6.add_run('que se solicita que se declare el divorcio respecto del matrimonio celebrado por ambas partes.')

    parrafo7=document.add_paragraph()
    parrafo7.add_run("Las materias acordadas son las siguientes: ")
 
    parrafo7=document.add_paragraph()
    parrafo7.add_run('ALIMENTOS MENORES:').bold=True
    if data_load['huboMediacion']=='Si':
        parrafo7.add_run(' En proceso de mediación RIT ')
        parrafo7.add_run(data_load['ritMediacion'])
        parrafo7.add_run(', seguida ante ')
        parrafo7.add_run('VACÍO').bold=True    #a futuro, agregar lista de lugares en los que puede llevarse a cabo la mediación.
        parrafo7.add_run(', las partes estipulan que ')
        if data_load['pagaAlimentosMenores']=='Patrocinado':
            parrafo7.add_run(data_load['honorificoPatrocinado'])
            parrafo7.add_run(' ')
            parrafo7.add_run(data_load['nombrePatrocinado'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoPaternoPatrocinado'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoMaternoPatrocinado'])
            parrafo7.add_run(' pagará la suma de $')
            parrafo7.add_run(data_load['cuantiaAlimentosMenores'])
            parrafo7.add_run(' pesos, equivalente a ')
            parrafo7.add_run(ingresoMinimo(data_load['cuantiaAlimentosMenores'],data_load['cuantiaIMR']))
            parrafo7.add_run('.')
        elif data_load['pagaAlimentosMenores']=='Contraparte':
            parrafo7.add_run(data_load['honorificoContraparte'])
            parrafo7.add_run(' ')
            parrafo7.add_run(data_load['nombreContraparte'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoPaternoContraparte'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoMaternoContraparte'])
            parrafo7.add_run(' pagará la suma de $')
            parrafo7.add_run(data_load['cuantiaAlimentosMenores'])
            parrafo7.add_run(' pesos, equivalente a ')
            parrafo7.add_run(ingresoMinimo(data_load['cuantiaAlimentosMenores'],data_load['cuantiaIMR']))
            parrafo7.add_run('.')
        elif data_load['pagaAlimentosMenores']=='No hay':
            parrafo7.add_run('no se pagarán alimentos menores, debido a')
            if data_load['hijosEntrePartes']==0:
                parrafo7.add_run(' que las partes no tienen hijos en común.')
            elif data_load['hijosMenoresDeEdad']==0:
                parrafo7.add_run(' que no hay hijos menores de edad nacidos de las partes.')
            else:
                parrafo7.add_run('que ')
                parrafo7.add_run('VACÍO').bold=True
    elif data_load['huboMediacion']=='No': 
        parrafo7.add_run(' Las partes estipulan que ')
        if data_load['pagaAlimentosMenores']=='No hay':
            parrafo7.add_run('no se pagarán alimentos menores, debido a')
            if data_load['hijosEntrePartes']==0:
                parrafo7.add_run(' que las partes no tienen hijos en común.')
            elif data_load['hijosMenoresDeEdad']==0:
                parrafo7.add_run(' que no hay hijos menores de edad nacidos de las partes.')
            else:
                parrafo7.add_run('que ')
                parrafo7.add_run('VACÍO').bold=True          
        else:
            if data_load['pagaAlimentosMenores']=='Patrocinado':
                parrafo7.add_run(data_load['honorificoPatrocinado'])
                parrafo7.add_run(' ')
                parrafo7.add_run(data_load['nombrePatrocinado'])
                parrafo7.add_run(" ")
                parrafo7.add_run(data_load['apellidoPaternoPatrocinado'])
                parrafo7.add_run(" ")
                parrafo7.add_run(data_load['apellidoMaternoPatrocinado'])
                parrafo7.add_run(' pagará la suma de $')
                parrafo7.add_run(data_load['cuantiaAlimentosMenores'])
                parrafo7.add_run(' pesos, equivalente a ')
                parrafo7.add_run(ingresoMinimo(data_load['cuantiaAlimentosMenores'],data_load['cuantiaIMR']))
                parrafo7.add_run('.')
            elif data_load['pagaAlimentosMenores']=='Contraparte':
                parrafo7.add_run(data_load['honorificoContraparte'])
                parrafo7.add_run(' ')
                parrafo7.add_run(data_load['nombreContraparte'])
                parrafo7.add_run(" ")
                parrafo7.add_run(data_load['apellidoPaternoContraparte'])
                parrafo7.add_run(" ")
                parrafo7.add_run(data_load['apellidoMaternoContraparte'])
                parrafo7.add_run(' pagará la suma de $')
                parrafo7.add_run(data_load['cuantiaAlimentosMenores'])
                parrafo7.add_run(' pesos, equivalente a ')
                parrafo7.add_run(ingresoMinimo(data_load['cuantiaAlimentosMenores'],data_load['cuantiaIMR']))
                parrafo7.add_run('.')

    parrafo7=document.add_paragraph()
    parrafo7.add_run('ALIMENTOS MAYORES:').bold=True        #preguntar si es común que se demanden alimentos mayores en juicios de DMA.        
    parrafo7=document.add_paragraph()
    parrafo7.add_run('No se han demandado alimentos entre los cónyuges.')
    
    
    parrafo7=document.add_paragraph()
    parrafo7.add_run('CUIDADO PERSONAL:').bold=True
    parrafo7=document.add_paragraph()
    if int(data_load['hijosEntrePartes'] )==0:
        parrafo7.add_run(' No corresponde por no haber hijos entre las partes.')
    elif int(data_load['hijosEntrePartes'] )>0:
        if data_load['hijosMenoresDeEdad']==0 and len(data_load['listaCuidadoPersonal']==0):
            parrafo7.add_run(' No corresponde por no haber hijos menores de edad.')
        elif len(data_load['listaCuidadoPersonal'])>0:
            for x in data_load['listaCuidadoPersonal']:
                parrafo7.add_run('En el caso de ')
                parrafo7.add_run(x[1:])
                parrafo7.add_run(', su cuidado personal corresponderá a ')

                if x[0]=='1':
                    parrafo7.add_run(data_load['honorificoPatrocinado'])
                    parrafo7.add_run(' ')
                    parrafo7.add_run(data_load['nombrePatrocinado'])
                    parrafo7.add_run(" ")
                    parrafo7.add_run(data_load['apellidoPaternoPatrocinado'])
                    parrafo7.add_run(" ")
                    parrafo7.add_run(data_load['apellidoMaternoPatrocinado'])
                    parrafo7.add_run('.')
                elif x[0]=='2':
                    parrafo7.add_run(data_load['honorificoContraparte'])
                    parrafo7.add_run(' ')
                    parrafo7.add_run(data_load['nombreContraparte'])
                    parrafo7.add_run(" ")
                    parrafo7.add_run(data_load['apellidoPaternoContraparte'])
                    parrafo7.add_run(" ")
                    parrafo7.add_run(data_load['apellidoMaternoContraparte'])
                    parrafo7.add_run('.') 
                parrafo7=document.add_paragraph()
                
    parrafo7=document.add_paragraph()
    parrafo7.add_run('PATRIA POTESTAD:').bold=True
    parrafo7=document.add_paragraph()
    if int(data_load['hijosEntrePartes'] )==0:
        parrafo7.add_run(' No corresponde por no haber hijos entre las partes.')
    elif int(data_load['hijosEntrePartes'] )>0:
        if data_load['hijosMenoresDeEdad']==0 and len(data_load['listaPatriaPotestad']==0):
            parrafo7.add_run(' No corresponde por no haber hijos menores de edad.')
        elif len(data_load['listaPatriaPotestad'])>0:
            for x in data_load['listaPatriaPotestad']:
                parrafo7.add_run('En el caso de ')
                parrafo7.add_run(x[1:])
                parrafo7.add_run(', su cuidado personal corresponderá a ')

                if x[0]=='1':
                    parrafo7.add_run(data_load['honorificoPatrocinado'])
                    parrafo7.add_run(' ')
                    parrafo7.add_run(data_load['nombrePatrocinado'])
                    parrafo7.add_run(" ")
                    parrafo7.add_run(data_load['apellidoPaternoPatrocinado'])
                    parrafo7.add_run(" ")
                    parrafo7.add_run(data_load['apellidoMaternoPatrocinado'])
                    parrafo7.add_run('.')
                elif x[0]=='2':
                    parrafo7.add_run(data_load['honorificoContraparte'])
                    parrafo7.add_run(' ')
                    parrafo7.add_run(data_load['nombreContraparte'])
                    parrafo7.add_run(" ")
                    parrafo7.add_run(data_load['apellidoPaternoContraparte'])
                    parrafo7.add_run(" ")
                    parrafo7.add_run(data_load['apellidoMaternoContraparte'])
                    parrafo7.add_run('.') 
                parrafo7=document.add_paragraph()              

    parrafo7=document.add_paragraph()
    parrafo7.add_run('RELACIÓN DIRECTA Y REGULAR:').bold=True
    parrafo7=document.add_paragraph()
    parrafo7.add_run('A este respecto se establece lo siguiente:')
    parrafo7.add_run(' VACÍO').bold=True
    
    parrafo7=document.add_paragraph()
    parrafo7.add_run('COMPENSACIÓN ECONÓMICA:').bold=True    
    parrafo7=document.add_paragraph()   
    if 'no' in data_load['pideCompensacionEconomica']:
        parrafo7.add_run('Habiendo sido informados sobre la institución de la compensación económica, ambas partes renuncian expresamente a ella')
    elif data_load['pideCompensacionEconomica']=='Patrocinado':
        parrafo7.add_run(' Las partes acuerdan que ')
        parrafo7.add_run(data_load['honorificoContraparte'])
        parrafo7.add_run(' ')
        parrafo7.add_run(data_load['nombreContraparte'])
        parrafo7.add_run(" ")
        parrafo7.add_run(data_load['apellidoPaternoContraparte'])
        parrafo7.add_run(" ")
        parrafo7.add_run(data_load['apellidoMaternoContraparte'])  
        parrafo7.add_run(' se obliga a pagar a ')  
        parrafo7.add_run(data_load['honorificoPatrocinado'])
        parrafo7.add_run(' ')
        parrafo7.add_run(data_load['nombrePatrocinado'])
        parrafo7.add_run(" ")
        parrafo7.add_run(data_load['apellidoPaternoPatrocinado'])
        parrafo7.add_run(" ")
        parrafo7.add_run(data_load['apellidoMaternoPatrocinado'])         
    elif data_load['pideCompensacionEconomica']=='Contraparte':
        parrafo7.add_run(' Las partes acuerdan que ')
        parrafo7.add_run(data_load['honorificoPatrocinado'])
        parrafo7.add_run(' ')
        parrafo7.add_run(data_load['nombrePatrocinado'])
        parrafo7.add_run(" ")
        parrafo7.add_run(data_load['apellidoPaternoPatrocinado'])
        parrafo7.add_run(" ")
        parrafo7.add_run(data_load['apellidoMaternoPatrocinado'])  
        parrafo7.add_run(' se obliga a pagar a ')  
        parrafo7.add_run(data_load['honorificoContraparte'])
        parrafo7.add_run(' ')
        parrafo7.add_run(data_load['nombreContraparte'])
        parrafo7.add_run(" ")
        parrafo7.add_run(data_load['apellidoPaternoContraparte'])
        parrafo7.add_run(" ")
        parrafo7.add_run(data_load['apellidoMaternoContraparte']) 
    parrafo7.add_run(', a título de compensación económica, el valor de $')
    parrafo7.add_run(data_load['valorCompensacion'])
    parrafo7.add_run('. Éste se hará efectivo')
    if data_load['modoPagoCompensacion'] == 'Bien Inmueble':
        parrafo7.add_run(', como pago equivalente a dicho monto, mediante la cesión del inmueble ')
        parrafo7.add_run(' ubicado en ')
        parrafo7.add_run(data_load['direccionInmueblePagoCompensacion'])
        parrafo7.add_run(' número ')
        parrafo7.add_run(data_load['numeroInmueblePagoCompensacion'])       
        parrafo7.add_run(' de propiedad de ')
        if data_load['pideCompensacionEconomica']=='Contraparte':
            parrafo7.add_run(data_load['honorificoPatrocinado'])
            parrafo7.add_run(' ')
            parrafo7.add_run(data_load['nombrePatrocinado'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoPaternoPatrocinado'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoMaternoPatrocinado']) 
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo7.add_run(data_load['honorificoContraparte'])
            parrafo7.add_run(' ')
            parrafo7.add_run(data_load['nombreContraparte'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoPaternoContraparte'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoMaternoContraparte']) 
        parrafo7.add_run('.')



    elif data_load['modoPagoCompensacion'] == 'Derechos sobre un inmueble':
        parrafo7.add_run(', como pago equivalente a dicho monto, mediante la cesión de los derechos que ')
        if data_load['pideCompensacionEconomica']=='Contraparte':
            parrafo7.add_run(data_load['honorificoPatrocinado'])
            parrafo7.add_run(' ')
            parrafo7.add_run(data_load['nombrePatrocinado'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoPaternoPatrocinado'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoMaternoPatrocinado']) 
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo7.add_run(data_load['honorificoContraparte'])
            parrafo7.add_run(' ')
            parrafo7.add_run(data_load['nombreContraparte'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoPaternoContraparte'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoMaternoContraparte']) 
        parrafo7.add_run(' tiene sobre el inmueble ubicado en ')
        parrafo7.add_run(data_load['direccionInmueblePagoCompensacion'])
        parrafo7.add_run(' número ')
        parrafo7.add_run(data_load['numeroInmueblePagoCompensacion'])       
        parrafo7.add_run('.') 
    elif data_load['modoPagoCompensacion'] == 'Derechos sobre inmueble sociedad conyugal': 
        parrafo7.add_run(', como pago equivalente a dicho monto, mediante la cesión de los derechos que a ')
        if data_load['pideCompensacionEconomica']=='Contraparte':
            parrafo7.add_run(data_load['honorificoPatrocinado'])
            parrafo7.add_run(' ')
            parrafo7.add_run(data_load['nombrePatrocinado'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoPaternoPatrocinado'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoMaternoPatrocinado']) 
        elif data_load['pideCompensacionEconomica']=='Patrocinado':
            parrafo7.add_run(data_load['honorificoContraparte'])
            parrafo7.add_run(' ')
            parrafo7.add_run(data_load['nombreContraparte'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoPaternoContraparte'])
            parrafo7.add_run(" ")
            parrafo7.add_run(data_load['apellidoMaternoContraparte']) 
        parrafo7.add_run(' corresponden sobre el inmueble propiedad de la sociedad conyugal ubicado en ')
        parrafo7.add_run(data_load['direccionInmueblePagoCompensacion'])
        parrafo7.add_run(' número ')
        parrafo7.add_run(data_load['numeroInmueblePagoCompensacion'])       
        parrafo7.add_run('.')     
    elif data_load['modoPagoCompensacion'] == 'Dinero':
        parrafo7.add_run(' mediante el pago de la suma de ')
        parrafo7.add_run(data_load['valorCompensacion'])
        parrafo7.add_run('.')
    elif data_load['modoPagoCompensacion'] == 'Otro': 
        parrafo7.add_run(' mediante')    
        parrafo7.add_run(' VACÍO.').bold==True

    parrafo7=document.add_paragraph()
    parrafo7.add_run('RÉGIMEN PATRIMONIAL:').bold=True   
    parrafo7=document.add_paragraph()
    parrafo7.add_run(' El régimen patrimonial pactado al momento de celebración del matrimonio de ambas partes fue el de ')
    parrafo7.add_run(data_load['regimenMatrimonial'].lower())
    if data_load['regimenMatrimonial'] == 'Separación total de bienes':
        parrafo7.add_run(', no existiendo por lo tanto bienes que liquidar.')
    else:
        parrafo7.add_run(', sin que existan bienes que liquidar.')
     
    parrafo7=document.add_paragraph()
    parrafo7.add_run('POR TANTO,').bold=True 

    parrafo7=document.add_paragraph()
    parrafo7.add_run('SÍRVASE SS.').bold=True     
    parrafo7.add_run(' calificar el presente acuerdo de completo y suficiente, autorizarlo y aprobarlo por regular de manera completa y suficiente las materias del artículo 21 de la ley 19.947.')

    
    document.save('acuerdoSimple'+data_load['apellidoPaternoPatrocinado']+data_load['apellidoPaternoContraparte']+'.docx')

    print("I am a zombie, ZOMBIE!!")
